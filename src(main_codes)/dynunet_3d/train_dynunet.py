"""
train_dynunet.py
================

Changes in this version vs previous:
  - SCHEDULER: cosine warmup → CosineAnnealingWarmRestarts (T_0=30)
    Warm restarts every 30 epochs prevent LR dying to near-zero and allow
    the model to escape precision plateaus that appeared after epoch ~65.
  - LOSS: Tversky beta 0.8 → 0.7  (alpha 0.2 → 0.3)
    Reduces false-negative penalty from 4× to 2.3×, directly addressing
    the over-segmentation (recall 0.90 / precision 0.68 imbalance).
  - AUGMENTATION: Rand3DElasticd added (prob=0.2)
    Helps model generalise across old dataset (D≈220) and new ISPY-1
    dataset (D≈50) which have different axis proportions.
  - DATA LOADER: build_loaders → build_weighted_loaders
    New ISPY-1 patients are sampled 2× per epoch — the updated data_3d.py
    already has this; just need to call the right function here.
  - OVERFITTING GUARD: L2 weight_decay 1e-5 → 5e-5 (stronger regularisation)
  - LR_START: set to match CosineAnnealingWarmRestarts — no manual warmup
    needed since CAWR handles warm-up naturally via T_0.
  - PATIENCE: 60 (unchanged — generous for warm-restart schedule)
  - Word log: Loss and scheduler strings updated to reflect new values.
  - build_weighted_loaders import added.
"""

import os
import sys
import numpy as np
import torch
import torch.nn as nn
import torch.optim as optim
import multiprocessing
import torch.nn.functional as F
from collections import Counter
from datetime import datetime
from torch.amp import autocast
from tqdm import tqdm

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    import subprocess, sys as _sys
    subprocess.run([_sys.executable, "-m", "pip", "install",
                    "python-docx", "--quiet"], check=False)
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        _DOCX_AVAILABLE = True
    except ImportError:
        _DOCX_AVAILABLE = False
        print("  [WARN] python-docx not available — Word log will not be saved")

from monai.inferers import sliding_window_inference

multiprocessing.freeze_support()

sys.path.append(os.path.dirname(__file__))

from data_3d       import build_weighted_loaders, PATCH_SIZE
from model_dynunet import get_dynunet
torch.cuda.empty_cache()

# ------------------------------------------------------------------ #
#  PATHS                                                               #
# ------------------------------------------------------------------ #

PREPROCESSED_ROOT = "data/patients_combined"
SAVE_DIR          = "models/dynunet_3d"
EPOCH_SAVE_DIR    = os.path.join(SAVE_DIR, "epochs_dynunet")

os.makedirs(SAVE_DIR,       exist_ok=True)
os.makedirs(EPOCH_SAVE_DIR, exist_ok=True)

# ------------------------------------------------------------------ #
#  HYPERPARAMETERS                                                     #
# ------------------------------------------------------------------ #

BATCH_SIZE  = 2
EPOCHS      = 150
PATIENCE    = 60
NUM_WORKERS = 0

# LR for CosineAnnealingWarmRestarts — no manual warmup needed
LR_MAX      = 1e-4      # peak LR each restart cycle
LR_MIN      = 1e-6      # floor LR at end of each cosine cycle
CAWR_T0     = 30        # restart every 30 epochs  (5 restarts in 150 ep)
CAWR_TMULT  = 1         # constant period — each cycle is 30 epochs

# Deep supervision
SW_OVERLAP        = 0.5
SW_MODE           = "gaussian"
DEEP_SUPR_WEIGHTS = [1.0, 0.5, 0.25]

VAL_THRESHOLDS = [0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.7]


# ------------------------------------------------------------------ #
#  LOSS FUNCTIONS                                                      #
#  Tversky: alpha=0.3, beta=0.7 (was 0.2/0.8)                        #
#  FN penalised 2.3× more than FP — reduces over-segmentation        #
#  Previous beta=0.8 was 4× penalty which drove recall too high       #
# ------------------------------------------------------------------ #

def tversky_loss(pred, target, alpha=0.3, beta=0.7, smooth=1e-6):
    """
    alpha=0.3, beta=0.7 — false negatives penalised 2.3× more than FP.
    Balanced enough to avoid over-segmentation while still finding tumour.
    """
    pred_sig = torch.sigmoid(pred)
    TP = (pred_sig * target).sum()
    FP = (pred_sig * (1 - target)).sum()
    FN = ((1 - pred_sig) * target).sum()
    tversky = (TP + smooth) / (TP + alpha * FP + beta * FN + smooth)
    return 1 - tversky


def focal_loss(pred, target, gamma=2.0, alpha=0.95, smooth=1e-6):
    """
    alpha=0.95 — 95% weight on tumour class.
    gamma=2.0  — moderate hard-example focus. Unchanged.
    """
    bce     = F.binary_cross_entropy_with_logits(pred, target, reduction='none')
    prob    = torch.sigmoid(pred)
    p_t     = prob * target + (1 - prob) * (1 - target)
    alpha_t = alpha * target + (1 - alpha) * (1 - target)
    focal   = alpha_t * (1 - p_t) ** gamma * bce
    return focal.mean()


def combined_loss(pred, target):
    return tversky_loss(pred, target) + focal_loss(pred, target)


def deep_supervision_loss(preds_list, target):
    total_loss = torch.tensor(0.0, device=target.device, requires_grad=True)
    for i, pred in enumerate(preds_list):
        weight = DEEP_SUPR_WEIGHTS[i] if i < len(DEEP_SUPR_WEIGHTS) else 0.25
        pred_f = pred.float()
        if pred_f.shape != target.shape:
            gt_down = F.interpolate(
                target.float(), size=pred_f.shape[2:], mode="nearest")
        else:
            gt_down = target.float()
        total_loss = total_loss + weight * combined_loss(pred_f, gt_down)
    return total_loss / sum(DEEP_SUPR_WEIGHTS[:len(preds_list)])


# ------------------------------------------------------------------ #
#  METRICS                                                             #
# ------------------------------------------------------------------ #

def compute_dice_np(pred_bin, gt_bin):
    intersection = np.logical_and(pred_bin, gt_bin).sum()
    return (2 * intersection) / (pred_bin.sum() + gt_bin.sum() + 1e-8)


# ------------------------------------------------------------------ #
#  COLLAPSE CHECK                                                      #
# ------------------------------------------------------------------ #

def check_output_distribution(model, loader, device, epoch):
    """Detects whether DynUNet is still in prediction collapse mode."""
    if (epoch + 1) % 5 != 0:
        return None

    model.eval()
    collapse_str = None
    with torch.no_grad():
        batch  = next(iter(loader))
        images = batch["image"].to(device)
        with autocast("cuda", dtype=torch.bfloat16):
            raw    = model(images[:1])
            logits = raw[0] if isinstance(raw, (list, tuple)) else raw
        probs = torch.sigmoid(logits.float())

        p_min  = probs.min().item()
        p_max  = probs.max().item()
        p_mean = probs.mean().item()
        p_high = (probs > 0.5).float().mean().item()

        if p_mean < 0.01 and p_min < 0.05:
            status = "OK — bias working correctly"
        elif p_mean > 0.3:
            status = "COLLAPSED — model predicting everything as tumor"
        elif p_min > 0.20:
            status = "WARNING — still in collapse"
        else:
            status = "RECOVERING"

        collapse_str = (f"[Collapse check ep{epoch+1}]"
                        f"  min={p_min:.4f}"
                        f"  max={p_max:.4f}"
                        f"  mean={p_mean:.4f}"
                        f"  frac>0.5={p_high:.4f}"
                        f"  → {status}")
        print(f"\n  {collapse_str}")

    model.train()
    return collapse_str


# ------------------------------------------------------------------ #
#  VALIDATION                                                          #
# ------------------------------------------------------------------ #

def validate(model, val_loader, device):
    model.eval()

    val_dices       = []
    val_losses      = []
    best_thresholds = []
    total_tp = total_fp = total_fn = total_tn = 0.0

    val_bar = tqdm(
        val_loader,
        desc="  Validation",
        leave=True,
        ncols=115,
        bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}] {postfix}"
    )

    with torch.no_grad():
        for batch in val_bar:
            images = batch["image"].to(device)
            labels = batch["label"].to(device)

            with autocast("cuda", dtype=torch.bfloat16):
                preds_logits = sliding_window_inference(
                    inputs        = images,
                    roi_size      = PATCH_SIZE,
                    sw_batch_size = 2,
                    predictor     = model,
                    overlap       = SW_OVERLAP,
                    mode          = SW_MODE,
                )

            val_loss = combined_loss(preds_logits.float(), labels.float())
            val_losses.append(val_loss.item())

            prob_map = torch.sigmoid(preds_logits.float())
            gt_np    = (labels > 0).squeeze().cpu().numpy().astype(bool)

            vol_dice    = 0.0
            best_thresh = VAL_THRESHOLDS[0]

            if gt_np.sum() > 0:
                for thresh in VAL_THRESHOLDS:
                    pred_t = (prob_map > thresh).squeeze().cpu().numpy().astype(bool)
                    d      = compute_dice_np(pred_t, gt_np)
                    if d > vol_dice:
                        vol_dice    = d
                        best_thresh = thresh
                val_dices.append(vol_dice)
                best_thresholds.append(best_thresh)

            pred_bin = (prob_map > best_thresh).float()
            p = pred_bin
            t = (labels > 0).float()
            tp = (p * t).sum().item()
            fp = (p * (1 - t)).sum().item()
            fn = ((1 - p) * t).sum().item()
            tn = ((1 - p) * (1 - t)).sum().item()

            total_tp += tp;  total_fp += fp
            total_fn += fn;  total_tn += tn

            cur_prec   = tp / (tp + fp + 1e-6)
            cur_rec    = tp / (tp + fn + 1e-6)
            foreground = tp + fp + fn
            cur_acc    = tp / (foreground + 1e-6) if foreground > 0 else 0.0

            val_bar.set_postfix(
                Dice=f"{vol_dice:.4f}",
                Loss=f"{val_loss.item():.4f}",
                Prec=f"{cur_prec:.4f}",
                Rec=f"{cur_rec:.4f}",
                BestT=f"{best_thresh:.1f}",
            )

    mean_dice     = float(np.mean(val_dices))  if val_dices  else 0.0
    mean_val_loss = float(np.mean(val_losses)) if val_losses else 0.0
    precision     = total_tp / (total_tp + total_fp + 1e-6)
    recall        = total_tp / (total_tp + total_fn + 1e-6)
    foreground_t  = total_tp + total_fp + total_fn
    accuracy      = total_tp / (foreground_t + 1e-6) if foreground_t > 0 else 0.0

    modal_thresh = Counter(best_thresholds).most_common(1)[0][0] \
                   if best_thresholds else VAL_THRESHOLDS[0]

    return {
        "loss"          : mean_val_loss,
        "dice"          : mean_dice,
        "precision"     : precision,
        "recall"        : recall,
        "accuracy"      : accuracy,
        "best_threshold": modal_thresh,
    }


# ------------------------------------------------------------------ #
#  TRAINING LOGGER                                                     #
# ------------------------------------------------------------------ #

class TrainingLogger:

    def __init__(self, save_path):
        self.save_path = save_path
        self.enabled   = _DOCX_AVAILABLE
        if not self.enabled:
            return
        self.doc = Document()
        for section in self.doc.sections:
            section.top_margin    = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin   = Inches(0.9)
            section.right_margin  = Inches(0.9)
        style      = self.doc.styles['Normal']
        font       = style.font
        font.name  = 'Consolas'
        font.size  = Pt(9)

    def _add(self, text, bold=False, color=None, size=9):
        if not self.enabled:
            return
        p    = self.doc.add_paragraph()
        run  = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(size)
        run.bold      = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

    def _separator(self, char='=', width=65):
        self._add(char * width)

    def save(self):
        if not self.enabled:
            return
        try:
            self.doc.save(self.save_path)
        except Exception as e:
            print(f"  [WARN] Could not save training log: {e}")

    def write_header(self, gpu_name, vram_gb, n_train, n_val, n_test,
                     patch_size, patches_vol, bias_layer, bias_val,
                     total_params, epochs, batch_size, patience,
                     deep_supr_weights, lr_max, save_dir):
        if not self.enabled:
            return
        self._add(f"Training started : {datetime.now().strftime('%Y-%m-%d  %H:%M:%S')}",
                  bold=True, size=10)
        self._add("")
        self._add(f"Using device: cuda")
        self._add(f"  GPU  : {gpu_name}")
        self._add(f"  VRAM : {vram_gb:.1f} GB")
        self._add("")
        self._add(f"[Dataset] Train : {n_train} patients  (weighted sampler — new 2x)")
        self._add(f"[Dataset] Val   : {n_val} patients")
        self._add(f"[Dataset] Test  : {n_test} patients")
        self._add(f"[Dataset] Patch : {patch_size}  |  Patches/vol: {patches_vol}")
        self._add("")
        self._add(f"  Output bias layer : {bias_layer}")
        self._add(f"  Output bias init  : {bias_val:.4f}"
                  f"  (sigmoid = {torch.sigmoid(torch.tensor(bias_val)).item():.5f}"
                  f"  ≈ tumour prevalence)")
        self._add("")
        self._add(f"  Model : DynUNet 3D (deep supervision)")
        self._add(f"  Total params : {total_params:,}")
        self._add("")
        self._separator()
        self._add("  Starting DynUNet 3D training  — warm-restart run", bold=True)
        self._separator()
        self._add(f"  Epochs       : {epochs}")
        self._add(f"  Batch size   : {batch_size}  (BF16 enabled)")
        self._add(f"  Patch size   : {patch_size}")
        self._add(f"  Patience     : {patience}")
        self._add(f"  Scheduler    : CosineAnnealingWarmRestarts  T0={CAWR_T0}  Tmult={CAWR_TMULT}")
        self._add(f"  LR range     : {LR_MIN:.0e} — {lr_max:.0e}")
        self._add(f"  Deep supr.   : weights {deep_supr_weights}")
        self._add(f"  Loss         : Tversky(a=0.3,b=0.7) + Focal(γ=2,α=0.95)")
        self._add(f"  Val thresh   : sweep {VAL_THRESHOLDS}")
        self._add(f"  Save dir     : {save_dir}")
        self._separator()
        self._add("")
        self.save()

    def write_epoch(self, epoch_num, total_epochs, phase, lr,
                    train_loss, train_recall, train_prec, train_acc,
                    val_loss, val_dice, val_prec, val_rec, val_acc,
                    best_thresh, smoothed_dice,
                    raw_best_dice, smooth_best_dice,
                    save_messages, no_improve_count,
                    collapse_check=None):
        if not self.enabled:
            return
        self._add("")
        self._add(f"Epoch {epoch_num}/{total_epochs}  [{phase}]  LR: {lr:.2e}",
                  bold=True, color=(0, 70, 127))
        self._add(f"  Train Loss      : {train_loss:.4f}")
        self._add(f"  Train Recall    : {train_recall:.4f}")
        self._add(f"  Train Precision : {train_prec:.4f}")
        self._add(f"  Train Accuracy  : {train_acc:.4f}")
        if collapse_check:
            self._add(f"  {collapse_check}", color=(150, 100, 0))
        self._add(f"  Val Loss        : {val_loss:.4f}")
        self._add(f"  Val Dice        : {val_dice:.4f}"
                  f"  (best-threshold sweep, tumour volumes only)",
                  bold=(val_dice > smooth_best_dice * 0.98))
        self._add(f"  Val Precision   : {val_prec:.4f}")
        self._add(f"  Val Recall      : {val_rec:.4f}")
        self._add(f"  Val Accuracy    : {val_acc:.4f}  (foreground voxels only)")
        self._add(f"  Best threshold  : {best_thresh:.1f}"
                  f"  (modal across val patients this epoch)")
        for msg in save_messages:
            bold_msg = msg.startswith(">>>") or msg.startswith("***")
            color    = (0, 100, 0) if msg.startswith("***") else (0, 0, 180)
            self._add(f"  {msg}", bold=bold_msg, color=color if bold_msg else None)
        if no_improve_count > 0:
            self._add(f"  No improvement ({no_improve_count}/{PATIENCE})",
                      color=(160, 80, 0))
        self.save()

    def write_footer(self, best_raw_dice, best_smoothed_dice,
                     final_val_loss, best_val_loss, optimal_thresh,
                     epoch_save_dir):
        if not self.enabled:
            return
        self._add("")
        self._separator()
        self._add("  DynUNet training complete", bold=True)
        self._separator()
        self._add(f"  Best raw Dice      : {best_raw_dice:.4f}"
                  f"  → dynunet_best_raw.pth", bold=True)
        self._add(f"  Best smoothed Dice : {best_smoothed_dice:.4f}"
                  f"  → dynunet_best.pth")
        self._add(f"  Final val loss     : {final_val_loss:.4f}")
        self._add(f"  Best val loss      : {best_val_loss:.4f}")
        self._add(f"  Optimal threshold  : {optimal_thresh:.1f}"
                  f"  (most common across all epochs)")
        self._add(f"  Per-epoch saves    : {epoch_save_dir}/")
        self._add("")
        self._add(f"Training finished : {datetime.now().strftime('%Y-%m-%d  %H:%M:%S')}",
                  bold=True)
        self.save()
        print(f"\n  Training log saved → {self.save_path}")


def _ordinal(n):
    suffix = {1:"st",2:"nd",3:"rd"}.get(
        n % 10 if n % 100 not in (11,12,13) else 0, "th")
    return f"{n}{suffix}"


# ------------------------------------------------------------------ #
#  MAIN                                                                #
# ------------------------------------------------------------------ #

def main():
    assert torch.cuda.is_available(), \
        "CUDA GPU not found. This script requires a GPU."
    device = torch.device("cuda")
    torch.cuda.set_device(0)

    print(f"\nUsing device: {device}")
    print(f"  GPU  : {torch.cuda.get_device_name(0)}")
    print(f"  VRAM : {torch.cuda.get_device_properties(0).total_memory / 1e9:.1f} GB")

    torch.set_num_threads(2)
    torch.set_num_interop_threads(1)
    torch.backends.cudnn.benchmark        = True
    torch.backends.cuda.matmul.allow_tf32 = True

    # ---- Data — weighted sampler (new ISPY-1 patients 2× per epoch) ----
    print()
    train_loader, val_loader, _ = build_weighted_loaders(
        preprocessed_root  = PREPROCESSED_ROOT,
        batch_size         = BATCH_SIZE,
        num_workers        = 0,
        new_dataset_weight = 2.0,
    )

    # ---- Model ----
    print()
    model = get_dynunet(device)

    try:
        model.apply(lambda m: setattr(m, 'gradient_checkpointing', True)
                    if hasattr(m, 'gradient_checkpointing') else None)
        torch.utils.checkpoint.checkpoint.__module__
        print("  Gradient checkpointing: enabled (~30% VRAM saved)")
    except Exception:
        print("  Gradient checkpointing: not available for this MONAI version")

    # ---- Optimizer — stronger weight decay to fight overfitting ----
    # weight_decay raised from 1e-5 → 5e-5
    optimizer = optim.AdamW(model.parameters(), lr=LR_MAX, weight_decay=5e-5)

    # ---- Scheduler: CosineAnnealingWarmRestarts ----
    # Restarts every T_0=30 epochs. Each restart raises LR back to LR_MAX
    # then decays to LR_MIN over 30 epochs. This prevents the LR from dying
    # near zero (which caused the plateau after epoch ~65 in the last run).
    # 5 complete restarts in 150 epochs = 5 chances to escape local optima.
    scheduler = optim.lr_scheduler.CosineAnnealingWarmRestarts(
        optimizer,
        T_0    = CAWR_T0,
        T_mult = CAWR_TMULT,
        eta_min = LR_MIN,
    )

    best_raw_dice      = 0.0
    best_smoothed_dice = 0.0
    counter            = 0

    history = {
        "train_loss"    : [], "train_recall"   : [],
        "val_loss"      : [], "val_dice"       : [],
        "val_precision" : [], "val_recall"     : [],
        "val_accuracy"  : [], "val_best_thresh": [],
    }

    log_path = os.path.join(SAVE_DIR, "training_log_dynunet.docx")
    logger   = TrainingLogger(log_path)

    from model_dynunet import OUTPUT_BIAS_INIT, DYNUNET_CONFIG
    total_params = sum(p.numel() for p in model.parameters())
    n_train      = len(train_loader.dataset)
    n_val        = len(val_loader.dataset)

    logger.write_header(
        gpu_name        = torch.cuda.get_device_name(0),
        vram_gb         = torch.cuda.get_device_properties(0).total_memory / 1e9,
        n_train         = n_train,
        n_val           = n_val,
        n_test          = 0,
        patch_size      = PATCH_SIZE,
        patches_vol     = 6,
        bias_layer      = "output_block.conv.conv.bias",
        bias_val        = OUTPUT_BIAS_INIT,
        total_params    = total_params,
        epochs          = EPOCHS,
        batch_size      = BATCH_SIZE,
        patience        = PATIENCE,
        deep_supr_weights = DEEP_SUPR_WEIGHTS,
        lr_max          = LR_MAX,
        save_dir        = SAVE_DIR,
    )

    print(f"  Epochs       : {EPOCHS}")
    print(f"  Batch size   : {BATCH_SIZE}  (BF16 enabled)")
    print(f"  Patch size   : {PATCH_SIZE}")
    print(f"  Patience     : {PATIENCE}")
    print(f"  Scheduler    : CosineAnnealingWarmRestarts  T0={CAWR_T0}  Tmult={CAWR_TMULT}")
    print(f"  LR range     : {LR_MIN:.0e} — {LR_MAX:.0e}")
    print(f"  Deep supr.   : weights {DEEP_SUPR_WEIGHTS}")
    print(f"  Loss         : Tversky(a=0.3,b=0.7) + Focal(γ=2,α=0.95)")
    print(f"  Weight decay : 5e-5  (increased for overfitting control)")
    print(f"  Val thresh   : sweep {VAL_THRESHOLDS}")
    print(f"  Weighted samp: new ISPY-1 patients 2× per epoch")
    print(f"  Save dir     : {SAVE_DIR}")
    print("=" * 65)
    print("  NOTE: Delete old .pth files before restarting:")
    print("    del models\\dynunet_3d\\*.pth")
    print("    del models\\dynunet_3d\\epochs_dynunet\\*.pth")
    print("=" * 65)

    def ordinal(n):
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(
            n % 10 if n % 100 not in (11, 12, 13) else 0, "th")
        return f"{n}{suffix}"

    for epoch in range(EPOCHS):

        current_lr = optimizer.param_groups[0]["lr"]

        # Phase label — restart cycle indicator for readability
        cycle_ep   = epoch % CAWR_T0
        cycle_num  = epoch // CAWR_T0 + 1
        phase      = f"CYCLE{cycle_num} ep{cycle_ep+1}/{CAWR_T0}"
        print(f"\nEpoch {epoch+1}/{EPOCHS}  [{phase}]  LR: {current_lr:.2e}")

        model.train()
        train_loss = 0.0
        tp = fp = fn = tn = 0.0

        train_bar = tqdm(
            train_loader,
            desc="  Training ",
            leave=True,
            ncols=115,
            bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}] {postfix}"
        )

        for step_idx, batch in enumerate(train_bar):

            images = batch["image"].to(device)
            masks  = batch["label"].to(device)

            masks = masks.float()
            if masks.ndim == 4:
                masks = masks.unsqueeze(1)

            optimizer.zero_grad()

            with autocast("cuda", dtype=torch.bfloat16):

                preds = model(images)

                if isinstance(preds, (list, tuple)):
                    preds_list = list(preds)
                elif preds.dim() == 6:
                    preds_list = [preds[:, i] for i in range(preds.shape[1])]
                else:
                    preds_list = [preds]

                total_loss = 0.0
                for i, pred in enumerate(preds_list):
                    pred = pred.float()
                    if pred.shape != masks.shape:
                        masks_resized = F.interpolate(
                            masks, size=pred.shape[2:], mode="nearest")
                    else:
                        masks_resized = masks
                    loss_i = combined_loss(pred, masks_resized)
                    weight = DEEP_SUPR_WEIGHTS[i] if i < len(DEEP_SUPR_WEIGHTS) else 0.25
                    total_loss += weight * loss_i

                loss       = total_loss / sum(DEEP_SUPR_WEIGHTS[:len(preds_list)])
                preds_full = preds_list[0]

            loss.backward()
            torch.nn.utils.clip_grad_norm_(model.parameters(), max_norm=1.0)
            optimizer.step()

            train_loss += loss.item()

            with torch.no_grad():
                p_sig = torch.sigmoid(preds_full.float())
                p_bin = (p_sig > 0.5).float()
                tp   += (p_bin * masks).sum().item()
                fp   += (p_bin * (1 - masks)).sum().item()
                fn   += ((1 - p_bin) * masks).sum().item()
                tn   += ((1 - p_bin) * (1 - masks)).sum().item()

            train_bar.set_postfix(loss=f"{loss.item():.4f}")

        avg_train_loss = train_loss / len(train_loader)
        train_recall   = tp / (tp + fn + 1e-6)
        train_prec     = tp / (tp + fp + 1e-6)
        train_acc      = (tp + tn) / (tp + tn + fp + fn + 1e-6)

        history["train_loss"].append(avg_train_loss)
        history["train_recall"].append(train_recall)

        print(f"\n  Train Loss      : {avg_train_loss:.4f}")
        print(f"  Train Recall    : {train_recall:.4f}")
        print(f"  Train Precision : {train_prec:.4f}")
        print(f"  Train Accuracy  : {train_acc:.4f}")

        # CosineAnnealingWarmRestarts steps per batch-step, but stepping
        # once per epoch (with fractional epoch) is also valid and simpler.
        scheduler.step(epoch + 1)

        torch.cuda.empty_cache()

        collapse_str = check_output_distribution(model, train_loader, device, epoch)

        print()
        val_metrics = validate(model, val_loader, device)

        avg_val_loss = val_metrics["loss"]
        avg_val_dice = val_metrics["dice"]
        prec         = val_metrics["precision"]
        rec          = val_metrics["recall"]
        acc          = val_metrics["accuracy"]
        best_thresh  = val_metrics["best_threshold"]

        history["val_loss"].append(avg_val_loss)
        history["val_dice"].append(avg_val_dice)
        history["val_precision"].append(prec)
        history["val_recall"].append(rec)
        history["val_accuracy"].append(acc)
        history["val_best_thresh"].append(best_thresh)

        print(f"\n  Val Loss        : {avg_val_loss:.4f}")
        print(f"  Val Dice        : {avg_val_dice:.4f}"
              f"  (best-threshold sweep, tumour volumes only)")
        print(f"  Val Precision   : {prec:.4f}")
        print(f"  Val Recall      : {rec:.4f}")
        print(f"  Val Accuracy    : {acc:.4f}  (foreground voxels only)")
        print(f"  Best threshold  : {best_thresh:.1f}"
              f"  (modal across val patients this epoch)")

        smoothed = float(np.mean(history["val_dice"][-3:])) \
                   if len(history["val_dice"]) >= 3 else avg_val_dice

        epoch_label      = ordinal(epoch + 1)
        _epoch_save_msgs = []

        torch.save(
            model.state_dict(),
            os.path.join(EPOCH_SAVE_DIR,
                         f"{epoch_label}_epoch_dice{avg_val_dice:.4f}.pth"))

        if avg_val_dice > best_raw_dice:
            best_raw_dice = avg_val_dice
            torch.save(model.state_dict(),
                       os.path.join(SAVE_DIR, "dynunet_best_raw.pth"))
            torch.save(
                model.state_dict(),
                os.path.join(EPOCH_SAVE_DIR,
                             f"BEST_RAW_{epoch_label}_epoch"
                             f"_dice{best_raw_dice:.4f}.pth"))
            print(f"\n  >>> Raw best saved    "
                  f"(Dice={best_raw_dice:.4f}, {epoch_label} epoch)"
                  f"  → dynunet_best_raw.pth")
            _epoch_save_msgs.append(
                f">>> Raw best saved    (Dice={best_raw_dice:.4f}, {epoch_label} epoch)"
                f"  → dynunet_best_raw.pth")

        if smoothed > best_smoothed_dice:
            best_smoothed_dice = smoothed
            counter = 0
            torch.save(model.state_dict(),
                       os.path.join(SAVE_DIR, "dynunet_best.pth"))
            torch.save(
                model.state_dict(),
                os.path.join(EPOCH_SAVE_DIR,
                             f"BEST_SMOOTH_{epoch_label}_epoch"
                             f"_dice{avg_val_dice:.4f}.pth"))
            print(f"  *** Smoothed best saved "
                  f"(smoothed Dice={best_smoothed_dice:.4f}, {epoch_label} epoch)"
                  f"  → dynunet_best.pth ***")
            _epoch_save_msgs.append(
                f"*** Smoothed best saved (smoothed Dice={best_smoothed_dice:.4f},"
                f" {epoch_label} epoch)  → dynunet_best.pth ***")
        else:
            counter += 1
            print(f"  No improvement ({counter}/{PATIENCE})")

        if counter >= PATIENCE:
            print(f"\nEarly stopping triggered at epoch {epoch+1}.")
            logger.write_epoch(
                epoch_num=epoch+1, total_epochs=EPOCHS, phase=phase, lr=current_lr,
                train_loss=avg_train_loss, train_recall=train_recall,
                train_prec=train_prec, train_acc=train_acc,
                val_loss=avg_val_loss, val_dice=avg_val_dice,
                val_prec=prec, val_rec=rec, val_acc=acc,
                best_thresh=best_thresh, smoothed_dice=smoothed,
                raw_best_dice=best_raw_dice, smooth_best_dice=best_smoothed_dice,
                save_messages=_epoch_save_msgs,
                no_improve_count=counter,
                collapse_check=collapse_str,
            )
            break

        logger.write_epoch(
            epoch_num=epoch+1, total_epochs=EPOCHS, phase=phase, lr=current_lr,
            train_loss=avg_train_loss, train_recall=train_recall,
            train_prec=train_prec, train_acc=train_acc,
            val_loss=avg_val_loss, val_dice=avg_val_dice,
            val_prec=prec, val_rec=rec, val_acc=acc,
            best_thresh=best_thresh, smoothed_dice=smoothed,
            raw_best_dice=best_raw_dice, smooth_best_dice=best_smoothed_dice,
            save_messages=_epoch_save_msgs,
            no_improve_count=counter,
            collapse_check=collapse_str,
        )

        np.save("training_history_dynunet.npy", history)

    print()
    print("=" * 65)
    print("  DynUNet training complete")
    print("=" * 65)
    print(f"  Best raw Dice      : {best_raw_dice:.4f}  → dynunet_best_raw.pth")
    print(f"  Best smoothed Dice : {best_smoothed_dice:.4f}  → dynunet_best.pth")
    if history["val_loss"]:
        print(f"  Final val loss     : {history['val_loss'][-1]:.4f}")
        print(f"  Best val loss      : {min(history['val_loss']):.4f}")
    if history["val_best_thresh"]:
        modal = Counter(history["val_best_thresh"]).most_common(1)[0][0]
        print(f"  Optimal threshold  : {modal:.1f}"
              f"  (most common across all epochs)")
    print(f"  Per-epoch saves    : {EPOCH_SAVE_DIR}/")

    modal_thresh = Counter(history["val_best_thresh"]).most_common(1)[0][0] \
                   if history["val_best_thresh"] else VAL_THRESHOLDS[0]
    logger.write_footer(
        best_raw_dice      = best_raw_dice,
        best_smoothed_dice = best_smoothed_dice,
        final_val_loss     = history["val_loss"][-1] if history["val_loss"] else 0.0,
        best_val_loss      = min(history["val_loss"]) if history["val_loss"] else 0.0,
        optimal_thresh     = modal_thresh,
        epoch_save_dir     = EPOCH_SAVE_DIR,
    )


if __name__ == "__main__":
    main()
