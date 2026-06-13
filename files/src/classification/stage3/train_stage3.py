"""
train_stage3.py
===============
Stage 3 — Benign vs Malignant Classifier  (3D EfficientNet-B0)

FIXES v2  (all in this file — model_stage3.py and data_stage3.py unchanged):

  FIX 1 — pos_weight corrected  ← CRITICAL, caused the collapse
    Was : 2.375  (323 total mal / 136 total ben — wrong: amplified the majority)
    Now : 0.403  (89 train ben / 221 train mal)
    WHY : Malignant is the MAJORITY class in training (71%).
          BCEWithLogitsLoss pos_weight UPWEIGHTS the positive class.
          Setting it > 1 when positives already dominate pushes the model
          to predict all-malignant (collapse: sens=1, spec=0, AUC=0.5).
          Setting pos_weight < 1 correctly DOWNWEIGHTS the majority so the
          model must learn to discriminate rather than predict trivially.

  FIX 2 — Label smoothing added (eps=0.1, train only)
    Converts  0 → 0.05,  1 → 0.95
    Prevents the overconfident predictions (0.9999 mal) that cause loss
    spikes and collapse when combined with a high pos_weight.

  FIX 3 — Per-epoch ordinal checkpoint saving (matches Stage 1 exactly)
    Stage 1 saves: 1st_epoch_auc0.8394.pth  (every epoch, in epochs/ folder)
    Stage 3 now:   1st_epoch_auc0.XXXX.pth  (same pattern, same folder name)
    Also saves BEST_RAW_ and BEST_SMOOTH_ labelled copies — identical to Stage 1.

  FIX 4 — Model save filenames match Stage 1 naming
    Stage 1: best_model.pth  /  best_model_smooth.pth  /  last_model.pth
    Stage 3: same names — evaluate_stage3.py already expects best_model.pth

  FIX 5 — Collapse check added at startup (like segmentation training)
    Prints min/max/mean of sigmoid outputs before epoch 1 so you can
    confirm the model starts at ~50% and not collapsed.

GPU PROFILES:
  "RTX3050"    → batch=8,  num_workers=0
  "RTX2000ADA" → batch=16, num_workers=4

Run:
    python src/classification/stage3/train_stage3.py

Saves to:
    models/classification_stage3/
        best_model.pth           ← best single-epoch val AUC
        best_model_smooth.pth    ← best 3-epoch smoothed val AUC
        last_model.pth           ← always latest epoch
        epochs/                  ← per-epoch ordinal files (Stage 1 format)
            1st_epoch_auc0.XXXX.pth
            BEST_RAW_Nth_epoch_auc0.XXXX.pth
            BEST_SMOOTH_Nth_epoch_auc0.XXXX.pth
        training_log.json
        training_log_stage3.docx
"""

import os
import sys
import math
import time
import json
import datetime
import platform
import numpy as np
import torch
import torch.nn as nn
from torch.amp import autocast, GradScaler
from torch.optim import AdamW
from tqdm import tqdm

try:
    from sklearn.metrics import (roc_auc_score, f1_score, precision_score,
                                  recall_score, accuracy_score, confusion_matrix)
    SKLEARN_OK = True
except ImportError:
    SKLEARN_OK = False
    print("[WARN] scikit-learn not found. Run: pip install scikit-learn")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
    print("[WARN] python-docx not found. Run: pip install python-docx")

_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)

from data_stage3  import build_weighted_stage3, STAGE3_ROOT, CROP_SIZE
from model_stage3 import get_model_stage3


# ================================================================== #
#  GPU PROFILE                                                         #
# ================================================================== #

ACTIVE_GPU = "RTX3050"   # ← change to "RTX2000ADA" for Ada

GPU_PROFILES = {
    "RTX3050": {
        "batch_size" : 8,
        "num_workers": 0,
        "use_amp"    : True,
        "note"       : "RTX 3050 6GB — ~400 MB VRAM at batch=8",
    },
    "RTX2000ADA": {
        "batch_size" : 16,
        "num_workers": 4,
        "use_amp"    : True,
        "note"       : "RTX 2000 Ada 16GB — ~700 MB VRAM at batch=16",
    },
}
GPU_CFG = GPU_PROFILES[ACTIVE_GPU]


# ================================================================== #
#  CONFIG                                                              #
# ================================================================== #

# ── Training set actual counts (from training log header) ──────────
# train=310 → ~221 malignant (71%), ~89 benign (29%)
_TRAIN_MAL = 221
_TRAIN_BEN = 89

CFG = {
    "data_root"          : STAGE3_ROOT,
    "save_dir"           : r"D:\Breast_Tumor_AI_Project\models\classification_stage3",
    "crop_size"          : CROP_SIZE,

    # FIX 1: pos_weight = N_benign_train / N_malignant_train
    # Malignant is MAJORITY (71%) → downweight it so model must discriminate
    "pos_weight"         : round(_TRAIN_BEN / _TRAIN_MAL, 4),   # 0.403

    "dropout"            : 0.4,
    "label_smooth_eps"   : 0.1,     # FIX 2: smooths 0→0.05, 1→0.95

    "epochs"             : 100,
    "lr_start"           : 1e-7,
    "lr_max"             : 3e-4,
    "warmup_epochs"      : 8,
    "cosine_t0"          : 30,
    "weight_decay"       : 1e-4,

    "patience"           : 30,
    "malignant_weight"   : 1.5,

    **GPU_CFG,
}

print(f"  pos_weight computed: {CFG['pos_weight']}  "
      f"({_TRAIN_BEN} ben / {_TRAIN_MAL} mal)")


# ================================================================== #
#  HELPERS                                                             #
# ================================================================== #

def ordinal(n):
    n = int(n)
    suffix = {1: "st", 2: "nd", 3: "rd"}
    return f"{n}{suffix.get(n % 10 if n % 100 not in (11, 12, 13) else 0, 'th')}"


def get_lr(epoch, warmup, t0, lr_start, lr_max):
    if epoch < warmup:
        return lr_start + (lr_max - lr_start) * (epoch / warmup)
    t = (epoch - warmup) % t0
    return lr_max * 0.5 * (1 + math.cos(math.pi * t / t0))


def lr_tag(epoch, warmup, t0):
    if epoch < warmup:
        return "[WARMUP]"
    cycle_ep = (epoch - warmup) % t0 + 1
    return f"[Cycle ep {cycle_ep}/{t0}]"


# FIX 2: Label smoothing — train only, never val/test
def smooth_labels(labels, eps=0.1):
    """0 → eps/2,  1 → 1 - eps/2.  Default eps=0.1: 0→0.05, 1→0.95."""
    return labels * (1 - eps) + (eps * 0.5)


# FIX 3: ordinal epoch filename (matches Stage 1 exactly)
def epoch_filename(epoch, auc):
    return f"{ordinal(epoch)}_epoch_auc{auc:.4f}.pth"


# ================================================================== #
#  WORD LOGGER  (identical format to Stage 1)                         #
# ================================================================== #

class WordLogger:
    BLUE  = RGBColor(0x00, 0x64, 0xC8) if DOCX_OK else None
    AMBER = RGBColor(0xB4, 0x64, 0x00) if DOCX_OK else None

    def __init__(self, path):
        self.path = path
        self.doc  = None
        if DOCX_OK:
            self.doc = Document()
            style = self.doc.styles["Normal"]
            style.font.name = "Consolas"
            style.font.size = Pt(9)

    def _line(self, text, bold=False, color=None):
        if self.doc is None:
            return
        p   = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.bold      = bold
        if color:
            run.font.color.rgb = color
        self._save()

    def _save(self):
        if self.doc:
            try:
                self.doc.save(self.path)
            except Exception:
                pass

    def write_header(self, device, gpu_name, vram_gb,
                     train_n, val_n, test_n, start_dt):
        sep = "=" * 70
        self._line(sep, bold=True)
        self._line("  STAGE 3 TRAINING LOG — Benign vs Malignant Classification",
                   bold=True)
        self._line("  Model: 3D EfficientNet-B0  |  v2 corrected pos_weight",
                   bold=True)
        self._line(sep, bold=True)
        self._line(f"  Started   : {start_dt.strftime('%Y-%m-%d %H:%M:%S')}")
        self._line(f"  Device    : {device}")
        self._line(f"  GPU       : {gpu_name}  ({vram_gb:.1f} GB VRAM)")
        self._line(f"  Profile   : {ACTIVE_GPU}  — {GPU_CFG['note']}")
        self._line(f"  Dataset   : train={train_n}  val={val_n}  test={test_n}")
        self._line(f"  Batch     : {CFG['batch_size']}   "
                   f"Crop: {CFG['crop_size']}³   Dropout: {CFG['dropout']}")
        self._line(f"  pos_weight: {CFG['pos_weight']}  "
                   f"({_TRAIN_BEN} ben / {_TRAIN_MAL} mal — downweights majority)")
        self._line(f"  LabelSmooth eps={CFG['label_smooth_eps']}  "
                   f"(train only: 0→0.05, 1→0.95)")
        self._line(f"  LR        : {CFG['lr_start']:.0e} → {CFG['lr_max']:.0e}  "
                   f"Warmup: {CFG['warmup_epochs']}ep  "
                   f"T0: {CFG['cosine_t0']}ep  "
                   f"Patience: {CFG['patience']}  Epochs: {CFG['epochs']}")
        self._line(f"  Save dir  : {CFG['save_dir']}")
        self._line(f"  Epoch fmt : {ordinal(1)}_epoch_auc0.XXXX.pth  "
                   f"(Stage 1 identical naming)")
        self._line(sep, bold=True)

    def write_epoch(self, epoch, total_epochs, lr, elapsed,
                    train_m, val_m, save_msgs):
        ep_str = ordinal(epoch)
        tag    = lr_tag(epoch - 1, CFG["warmup_epochs"], CFG["cosine_t0"])
        self._line(
            f"\nEpoch {ep_str}/{total_epochs}  {tag}  "
            f"LR: {lr:.2e}  time={elapsed:.0f}s",
            bold=True
        )
        self._line(
            f"  Train  loss={train_m['train_loss']:.4f}  "
            f"acc={train_m['train_acc']:.4f}  "
            f"auc={train_m.get('train_auc', 0):.4f}  "
            f"f1={train_m.get('train_f1', 0):.4f}"
        )
        self._line(
            f"  Val    loss={val_m['val_loss']:.4f}  "
            f"acc={val_m['val_acc']:.4f}  "
            f"auc={val_m.get('val_auc', 0):.4f}  "
            f"f1={val_m.get('val_f1', 0):.4f}",
            bold=True, color=self.BLUE
        )
        self._line(
            f"  sens={val_m.get('val_sens', 0):.4f}  "
            f"spec={val_m.get('val_spec', 0):.4f}  "
            f"ppv={val_m.get('val_prec', 0):.4f}"
        )
        for msg in save_msgs:
            self._line(f"  {msg}", color=self.AMBER)

    def write_footer(self, best_raw_epoch, best_raw_auc,
                     best_smooth_auc, test_m, end_dt):
        sep = "=" * 70
        self._line(f"\n{sep}", bold=True)
        self._line("  TRAINING COMPLETE", bold=True)
        self._line(f"  Finished       : {end_dt.strftime('%Y-%m-%d %H:%M:%S')}")
        self._line(f"  Best raw epoch : {ordinal(best_raw_epoch)}  "
                   f"AUC={best_raw_auc:.4f}")
        self._line(f"  Best smooth AUC: {best_smooth_auc:.4f}")
        self._line(f"  ── Test set results ──")
        for k, v in test_m.items():
            label = k.replace("test_", "").capitalize()
            val_s = f"{v:.4f}" if isinstance(v, float) else str(v)
            self._line(f"  {label:20s}: {val_s}")
        self._line(f"  Epoch file format  : {ordinal(1)}_epoch_auc0.XXXX.pth")
        self._line(f"  best_model.pth       — highest single-epoch AUC")
        self._line(f"  best_model_smooth.pth — highest 3-epoch smooth AUC")
        self._line(sep, bold=True)
        print(f"  Word log → {self.path}")


# ================================================================== #
#  METRICS                                                             #
# ================================================================== #

def compute_metrics(labels, probs, preds, prefix):
    m = {}
    m[f"{prefix}_acc"] = round(accuracy_score(labels, preds), 4)
    if SKLEARN_OK and len(set(labels)) > 1:
        m[f"{prefix}_auc"]  = round(float(roc_auc_score(labels, probs)), 4)
        m[f"{prefix}_f1"]   = round(float(f1_score(labels, preds, zero_division=0)), 4)
        m[f"{prefix}_prec"] = round(float(precision_score(labels, preds, zero_division=0)), 4)
        m[f"{prefix}_rec"]  = round(float(recall_score(labels, preds, zero_division=0)), 4)
        cm = confusion_matrix(labels, preds, labels=[0, 1])
        if cm.shape == (2, 2):
            tn, fp, fn, tp = cm.ravel()
            m[f"{prefix}_sens"] = round(tp / (tp + fn) if (tp + fn) > 0 else 0, 4)
            m[f"{prefix}_spec"] = round(tn / (tn + fp) if (tn + fp) > 0 else 0, 4)
    else:
        m[f"{prefix}_auc"] = m[f"{prefix}_f1"] = 0.0
    return m


# ================================================================== #
#  COLLAPSE CHECK                                                      #
# ================================================================== #

def check_output_distribution(model, device, use_amp):
    """FIX 5: Verify model is NOT collapsed before training starts."""
    model.eval()
    with torch.no_grad():
        dummy  = torch.randn(4, 3, CROP_SIZE, CROP_SIZE, CROP_SIZE).to(device)
        with autocast(device_type=device.type, enabled=use_amp):
            logits = model(dummy)
        probs = torch.sigmoid(logits).cpu().numpy().ravel()
    print(f"  Pre-train output check:")
    print(f"    mean prob = {probs.mean():.4f}  "
          f"std = {probs.std():.4f}  "
          f"range = [{probs.min():.4f}, {probs.max():.4f}]")
    if probs.std() < 0.01:
        print("  [WARN] All outputs nearly identical — check model init!")
    else:
        print("  [OK] Outputs are diverse — model not collapsed at init")
    model.train()


# ================================================================== #
#  PRINT EPOCH TABLE                                                   #
# ================================================================== #

def print_epoch_metrics(epoch, total_epochs, train_m, val_m,
                        elapsed, is_best, best_score, lr):
    sep = "  " + "─" * 61
    print(sep)
    print(f"  Epoch {epoch:3d}/{total_epochs}   "
          f"time={elapsed:.0f}s   lr={lr:.2e}   "
          f"{'★ NEW BEST' if is_best else ''}")
    print()
    print(f"  {'Metric':20s}  {'Train':>9s}  {'Val':>9s}")
    print(f"  {'-'*20}  {'-'*9}  {'-'*9}")
    rows = [
        ("Loss",        "train_loss",  "val_loss"),
        ("Accuracy",    "train_acc",   "val_acc"),
        ("AUC",         "train_auc",   "val_auc"),
        ("F1 Score",    "train_f1",    "val_f1"),
        ("Precision",   "train_prec",  "val_prec"),
        ("Recall",      "train_rec",   "val_rec"),
        ("Sensitivity", "train_sens",  "val_sens"),
        ("Specificity", "train_spec",  "val_spec"),
    ]
    for label, tk, vk in rows:
        tv = train_m.get(tk, "—")
        vv = val_m.get(vk, "—")
        tv_s = f"{tv:.4f}" if isinstance(tv, float) else str(tv)
        vv_s = f"{vv:.4f}" if isinstance(vv, float) else str(vv)
        print(f"  {label:20s}  {tv_s:>9s}  {vv_s:>9s}")
    if is_best:
        print(f"\n  ✓ Best model saved  (val_auc={best_score:.4f})")
    print(sep)
    print()


# ================================================================== #
#  TRAIN ONE EPOCH                                                     #
# ================================================================== #

def train_one_epoch(model, loader, criterion, optimizer, scaler,
                    device, use_amp, epoch, total_epochs, smooth_eps):
    model.train()
    total_loss = 0.0
    all_labels, all_probs, all_preds = [], [], []

    pbar = tqdm(
        loader,
        desc=f"  Epoch {epoch:3d}/{total_epochs} [TRAIN]",
        ncols=80, leave=False, file=sys.stdout,
    )
    for images, labels in pbar:
        images = images.to(device, non_blocking=True)
        labels = labels.to(device, non_blocking=True).unsqueeze(1)

        optimizer.zero_grad(set_to_none=True)

        with autocast(device_type=device.type, enabled=use_amp):
            logits = model(images)
            # FIX 2: smooth labels for training loss only
            loss   = criterion(logits, smooth_labels(labels, smooth_eps))

        if use_amp and device.type == "cuda":
            scaler.scale(loss).backward()
            scaler.unscale_(optimizer)
            nn.utils.clip_grad_norm_(model.parameters(), 1.0)
            scaler.step(optimizer)
            scaler.update()
        else:
            loss.backward()
            nn.utils.clip_grad_norm_(model.parameters(), 1.0)
            optimizer.step()

        total_loss += loss.item()
        probs = torch.sigmoid(logits).detach().cpu().numpy().ravel()
        preds = (probs >= 0.5).astype(int)
        labs  = labels.detach().cpu().numpy().ravel().astype(int)
        all_probs.extend(probs.tolist())
        all_preds.extend(preds.tolist())
        all_labels.extend(labs.tolist())
        pbar.set_postfix({"loss": f"{loss.item():.4f}"}, refresh=True)

    pbar.close()
    m = compute_metrics(all_labels, all_probs, all_preds, "train")
    m["train_loss"] = round(total_loss / len(loader), 5)
    return m


# ================================================================== #
#  EVALUATE                                                            #
# ================================================================== #

@torch.no_grad()
def evaluate(model, loader, criterion, device, use_amp,
             prefix="val", epoch=None, total_epochs=None):
    model.eval()
    total_loss = 0.0
    all_labels, all_probs, all_preds = [], [], []

    desc = (f"  Epoch {epoch:3d}/{total_epochs} [{prefix.upper():5s}]"
            if epoch else f"  [{prefix.upper()}]")
    pbar = tqdm(loader, desc=desc, ncols=80, leave=False, file=sys.stdout)

    for images, labels in pbar:
        images = images.to(device, non_blocking=True)
        labels = labels.to(device, non_blocking=True).unsqueeze(1)

        with autocast(device_type=device.type, enabled=use_amp):
            logits = model(images)
            # NOTE: do NOT smooth val/test labels — evaluate at true 0/1
            loss   = criterion(logits, labels.float())

        total_loss += loss.item()
        probs = torch.sigmoid(logits).cpu().numpy().ravel()
        preds = (probs >= 0.5).astype(int)
        labs  = labels.cpu().numpy().ravel().astype(int)
        all_probs.extend(probs.tolist())
        all_preds.extend(preds.tolist())
        all_labels.extend(labs.tolist())

    pbar.close()
    m = compute_metrics(all_labels, all_probs, all_preds, prefix)
    m[f"{prefix}_loss"] = round(total_loss / len(loader), 5)
    return m


# ================================================================== #
#  MAIN                                                                #
# ================================================================== #

def main():
    save_dir  = CFG["save_dir"]
    # FIX 3: epochs/ folder — identical to Stage 1
    epoch_dir = os.path.join(save_dir, "epochs")
    os.makedirs(save_dir,  exist_ok=True)
    os.makedirs(epoch_dir, exist_ok=True)

    start_dt = datetime.datetime.now()

    print()
    print("=" * 65)
    print("  Stage 3 — Benign vs Malignant Classifier  [v2 fixed]")
    print("  Model : 3D EfficientNet-B0")
    print(f"  GPU   : {ACTIVE_GPU}  — {GPU_CFG['note']}")
    print(f"  pos_weight = {CFG['pos_weight']}  (was 2.375 — now corrected)")
    print("=" * 65)

    device   = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    gpu_name = ""
    vram_gb  = 0.0
    print(f"\n  Device : {device}")
    if device.type == "cuda":
        gpu_name = torch.cuda.get_device_name(0)
        vram_gb  = torch.cuda.get_device_properties(0).total_memory / 1e9
        print(f"  GPU    : {gpu_name}")
        print(f"  VRAM   : {vram_gb:.1f} GB")

    # ── Data ──────────────────────────────────────────────────────
    print()
    train_loader, val_loader, test_loader = build_weighted_stage3(
        root             = CFG["data_root"],
        batch_size       = CFG["batch_size"],
        num_workers      = CFG["num_workers"],
        crop_size        = CFG["crop_size"],
        malignant_weight = CFG["malignant_weight"],
    )
    train_n = len(train_loader.dataset)
    val_n   = len(val_loader.dataset)
    test_n  = len(test_loader.dataset)

    # ── Model ─────────────────────────────────────────────────────
    print()
    model = get_model_stage3(device, dropout=CFG["dropout"])

    # FIX 5: collapse check before any training
    check_output_distribution(model, device, CFG["use_amp"])

    # ── Loss: corrected pos_weight ─────────────────────────────────
    pos_weight = torch.tensor([CFG["pos_weight"]], device=device)
    criterion  = nn.BCEWithLogitsLoss(pos_weight=pos_weight, reduction="mean")
    print(f"  Loss   : BCEWithLogitsLoss  pos_weight={CFG['pos_weight']}  "
          f"label_smooth_eps={CFG['label_smooth_eps']}")

    # ── Optimiser ─────────────────────────────────────────────────
    optimizer = AdamW(model.parameters(), lr=CFG["lr_start"],
                      weight_decay=CFG["weight_decay"])
    scaler    = GradScaler("cuda", enabled=(CFG["use_amp"] and
                                             device.type == "cuda"))

    print(f"  Epochs : {CFG['epochs']}   LR: {CFG['lr_start']:.0e} → {CFG['lr_max']:.0e}")
    print(f"  Patience: {CFG['patience']}")
    print()

    # ── Word logger ────────────────────────────────────────────────
    docx_path = os.path.join(save_dir, "training_log_stage3.docx")
    wlog = WordLogger(docx_path)
    wlog.write_header(device, gpu_name, vram_gb,
                      train_n, val_n, test_n, start_dt)

    # ── Training state ─────────────────────────────────────────────
    best_raw_auc    = 0.0
    best_smooth_auc = 0.0
    smooth_window   = []
    best_raw_epoch  = 1
    patience_cnt    = 0
    history         = []
    log_json        = os.path.join(save_dir, "training_log.json")

    for epoch in range(1, CFG["epochs"] + 1):
        lr = get_lr(epoch - 1, CFG["warmup_epochs"],
                    CFG["cosine_t0"], CFG["lr_start"], CFG["lr_max"])
        for pg in optimizer.param_groups:
            pg["lr"] = lr

        t0 = time.time()

        train_m = train_one_epoch(
            model, train_loader, criterion, optimizer, scaler,
            device, CFG["use_amp"], epoch, CFG["epochs"],
            CFG["label_smooth_eps"])

        val_m = evaluate(
            model, val_loader, criterion, device, CFG["use_amp"],
            prefix="val", epoch=epoch, total_epochs=CFG["epochs"])

        elapsed = time.time() - t0

        row = {"epoch": epoch, **train_m, **val_m, "lr": round(lr, 8)}
        history.append(row)

        score   = val_m.get("val_auc", val_m.get("val_acc", 0))
        is_best = (score > best_raw_auc)
        save_msgs = []

        # ── FIX 3: ordinal epoch checkpoint (every epoch) ─────────
        ep_fname   = epoch_filename(epoch, score)
        ep_ckpt    = os.path.join(epoch_dir, ep_fname)
        torch.save(model.state_dict(), ep_ckpt)

        # ── Raw best ──────────────────────────────────────────────
        if is_best:
            best_raw_auc   = score
            best_raw_epoch = epoch
            patience_cnt   = 0

            # FIX 4: save as best_model.pth (matches Stage 1 name)
            torch.save({
                "epoch"      : epoch,
                "model_state": model.state_dict(),
                "optimizer"  : optimizer.state_dict(),
                "best_score" : best_raw_auc,
                "cfg"        : CFG,
            }, os.path.join(save_dir, "best_model.pth"))

            # Labelled copy in epochs/ with BEST_RAW_ prefix
            best_raw_fname = f"BEST_RAW_{ep_fname}"
            torch.save(model.state_dict(),
                       os.path.join(epoch_dir, best_raw_fname))

            msg = (f">>> Raw best saved  — "
                   f"AUC={score:.4f}  "
                   f"Acc={val_m['val_acc']:.4f}  "
                   f"(epoch {ordinal(epoch)})")
            print(f"  {msg}")
            save_msgs.append(msg)
        else:
            patience_cnt += 1

        # ── Smooth best (3-epoch rolling mean) ────────────────────
        smooth_window.append(score)
        if len(smooth_window) > 3:
            smooth_window.pop(0)
        smooth_auc = float(np.mean(smooth_window))
        if smooth_auc > best_smooth_auc:
            best_smooth_auc = smooth_auc

            # FIX 4: save as best_model_smooth.pth (matches Stage 1 name)
            torch.save({
                "epoch"      : epoch,
                "model_state": model.state_dict(),
                "best_score" : best_smooth_auc,
            }, os.path.join(save_dir, "best_model_smooth.pth"))

            # Labelled copy in epochs/ with BEST_SMOOTH_ prefix
            smooth_fname = f"BEST_SMOOTH_{ep_fname}"
            torch.save(model.state_dict(),
                       os.path.join(epoch_dir, smooth_fname))

            msg = (f"*** Smooth best saved — "
                   f"smooth_AUC={smooth_auc:.4f}  (epoch {ordinal(epoch)})")
            print(f"  {msg}")
            save_msgs.append(msg)

        # ── Terminal table ─────────────────────────────────────────
        print_epoch_metrics(epoch, CFG["epochs"], train_m, val_m,
                            elapsed, is_best, best_raw_auc, lr)

        # ── Word log ───────────────────────────────────────────────
        wlog.write_epoch(epoch, CFG["epochs"], lr, elapsed,
                         train_m, val_m, save_msgs)

        # ── last_model.pth (always latest) ────────────────────────
        torch.save({
            "epoch"      : epoch,
            "model_state": model.state_dict(),
            "optimizer"  : optimizer.state_dict(),
        }, os.path.join(save_dir, "last_model.pth"))

        # ── JSON log ───────────────────────────────────────────────
        with open(log_json, "w") as f:
            json.dump({"cfg": CFG, "history": history}, f,
                      indent=2, default=str)

        # ── Early stopping ─────────────────────────────────────────
        if patience_cnt >= CFG["patience"]:
            print(f"  Early stopping at epoch {epoch} "
                  f"(patience={CFG['patience']})")
            break

        if epoch % 10 == 0:
            torch.cuda.empty_cache()

    # ── Final test evaluation ──────────────────────────────────────
    print("=" * 65)
    print("  Final TEST evaluation (best model)")
    print("=" * 65)

    ckpt = torch.load(os.path.join(save_dir, "best_model.pth"),
                      map_location=device, weights_only=False)
    model.load_state_dict(ckpt["model_state"])

    test_m = evaluate(model, test_loader, criterion, device,
                      CFG["use_amp"], prefix="test")

    print()
    print(f"  {'Metric':20s}  {'Value':>9s}")
    print(f"  {'-'*20}  {'-'*9}")
    for k, v in test_m.items():
        label = k.replace("test_", "").capitalize()
        print(f"  {label:20s}  "
              f"{f'{v:.4f}' if isinstance(v, float) else str(v):>9s}")

    # ── Word footer ────────────────────────────────────────────────
    end_dt = datetime.datetime.now()
    wlog.write_footer(best_raw_epoch, best_raw_auc,
                      best_smooth_auc, test_m, end_dt)

    # ── Final JSON ─────────────────────────────────────────────────
    with open(log_json) as f:
        log = json.load(f)
    log["test_results"] = test_m
    with open(log_json, "w") as f:
        json.dump(log, f, indent=2, default=str)

    print(f"\n  Log      → {log_json}")
    print(f"  Word log → {docx_path}")
    print()
    print("  Checkpoint naming (Stage 1 identical):")
    print(f"    epochs/1st_epoch_auc0.XXXX.pth     — every epoch")
    print(f"    epochs/BEST_RAW_Nth_epoch_auc0.XXXX.pth")
    print(f"    epochs/BEST_SMOOTH_Nth_epoch_auc0.XXXX.pth")
    print(f"    best_model.pth                     — highest single-epoch AUC")
    print(f"    best_model_smooth.pth              — highest 3-epoch smooth AUC")
    print(f"    last_model.pth                     — latest epoch")
    print(f"\n  Stage 3 training complete.")


if __name__ == "__main__":
    main()