"""
train_3d.py — upgraded MONAI training for combined dataset
===========================================================

Key upgrades from previous version:
  - NEW: Cosine warm-restart LR (CosineAnnealingWarmRestarts, T_0=30)
    replaces single cosine decay. Restarts every 30 epochs so the model
    can escape the plateau seen from epoch 65+ in the previous run.
  - NEW: Rebalanced loss — Tversky alpha=0.3, beta=0.7 (was 0.2/0.8).
    Collapse is fully solved; balanced loss reduces over-segmentation
    and improves precision on the new-dataset patients.
  - NEW: build_weighted_loaders() — new ISPY-1 patients (IDs >100) are
    sampled 2x more often to close the geometry gap.
  - NEW: Word .docx training log (same as DynUNet logger).
  - BF16 autocast retained. GradScaler removed (caused zero-loss bug
    in BF16; BF16 has sufficient dynamic range without scaler).
  - PATIENCE raised to 60 (more epochs before early stopping).
  - history saves val_best_thresh per epoch.
  - Collapse check every 5 epochs retained.

Usage:
    python src/segmentation_3d/train_3d.py

Delete old checkpoints before restarting:
    del models\\segmentation_3d\\*.pth
    del models\\segmentation_3d\\epochs\\*.pth
"""

import os
import sys
import numpy as np
import torch
import torch.nn as nn
import torch.optim as optim
import multiprocessing
import torch.nn.functional as F
from collections import Counter
from datetime import datetime
from torch.amp import autocast
from tqdm import tqdm

# python-docx — auto-install if missing
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    _DOCX_AVAILABLE = True
except ImportError:
    import subprocess, sys as _sys
    subprocess.run([_sys.executable, "-m", "pip", "install",
                    "python-docx", "--quiet"], check=False)
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        _DOCX_AVAILABLE = True
    except ImportError:
        _DOCX_AVAILABLE = False
        print("  [WARN] python-docx not available — Word log disabled")

from monai.inferers import sliding_window_inference

multiprocessing.freeze_support()

sys.path.append(os.path.dirname(__file__))

from data_3d  import build_weighted_loaders, PATCH_SIZE
from model_3d import get_model

torch.cuda.empty_cache()

# ------------------------------------------------------------------ #
#  PATHS                                                               #
# ------------------------------------------------------------------ #

PREPROCESSED_ROOT = "data/patients_combined"
SAVE_DIR          = "models/segmentation_3d"
EPOCH_SAVE_DIR    = os.path.join(SAVE_DIR, "epochs")

os.makedirs(SAVE_DIR,       exist_ok=True)
os.makedirs(EPOCH_SAVE_DIR, exist_ok=True)

# ------------------------------------------------------------------ #
#  HYPERPARAMETERS                                                     #
# ------------------------------------------------------------------ #

BATCH_SIZE          = 4
EPOCHS              = 150
WARMUP_EPOCHS       = 10       # linear warmup before cosine restarts begin
LR_MAX              = 1e-4
LR_START            = 1e-8
PATIENCE            = 60
NUM_WORKERS         = 0        # must be 0 on Windows
NEW_DATASET_WEIGHT  = 2.0      # new patients sampled 2x per epoch

SW_OVERLAP    = 0.25
SW_MODE       = "gaussian"

# Cosine warm-restart period (epochs per cycle after warmup)
COSINE_T0     = 30    # restart every 30 epochs
COSINE_TMULT  = 1     # each restart same length (set 2 to double each time)

VAL_THRESHOLDS = [0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.7]


# ------------------------------------------------------------------ #
#  LOSS FUNCTIONS                                                      #
#  Rebalanced: alpha=0.3, beta=0.7                                    #
#  Previous run used 0.2/0.8 — still recall-heavy, causing            #
#  over-segmentation on new-dataset patients. 0.3/0.7 improves        #
#  precision while keeping recall high.                                #
# ------------------------------------------------------------------ #

def tversky_loss(pred, target, alpha=0.3, beta=0.7, smooth=1e-6):
    """
    alpha=0.3, beta=0.7 — FN penalised ~2.3x more than FP.
    Balanced now that collapse is solved. Reduces over-segmentation.
    """
    pred_sig = torch.sigmoid(pred)
    TP = (pred_sig * target).sum()
    FP = (pred_sig * (1 - target)).sum()
    FN = ((1 - pred_sig) * target).sum()
    tversky = (TP + smooth) / (TP + alpha * FP + beta * FN + smooth)
    return 1 - tversky


def focal_loss(pred, target, gamma=2.0, alpha=0.95, smooth=1e-6):
    """
    alpha=0.95 — 95% weight on tumour class.
    gamma=2.0  — moderate hard-example focus.
    """
    bce     = F.binary_cross_entropy_with_logits(pred, target, reduction='none')
    prob    = torch.sigmoid(pred)
    p_t     = prob * target + (1 - prob) * (1 - target)
    alpha_t = alpha * target + (1 - alpha) * (1 - target)
    focal   = alpha_t * (1 - p_t) ** gamma * bce
    return focal.mean()


def combined_loss(pred, target):
    """Tversky(0.3,0.7) + Focal(γ=2,α=0.95)"""
    return tversky_loss(pred, target) + focal_loss(pred, target)


# ------------------------------------------------------------------ #
#  METRICS                                                             #
# ------------------------------------------------------------------ #

def compute_dice_np(pred_bin, gt_bin):
    intersection = np.logical_and(pred_bin, gt_bin).sum()
    return (2 * intersection) / (pred_bin.sum() + gt_bin.sum() + 1e-8)


# ------------------------------------------------------------------ #
#  TRAINING LOGGER (Word .docx)                                        #
# ------------------------------------------------------------------ #

class TrainingLogger:
    """Writes epoch-by-epoch training log to a .docx file."""

    def __init__(self, save_path):
        self.save_path = save_path
        self.enabled   = _DOCX_AVAILABLE
        if not self.enabled:
            return
        self.doc = Document()
        for section in self.doc.sections:
            section.top_margin    = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin   = Inches(0.9)
            section.right_margin  = Inches(0.9)
        style = self.doc.styles['Normal']
        style.font.name = 'Consolas'
        style.font.size = Pt(9)

    def _add(self, text, bold=False, color=None, size=9):
        if not self.enabled:
            return
        p   = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

    def _sep(self, w=62):
        self._add('=' * w)

    def save(self):
        if not self.enabled:
            return
        try:
            self.doc.save(self.save_path)
        except Exception as e:
            print(f"  [WARN] Could not save training log: {e}")

    def write_header(self, gpu_name, vram_gb, n_train, n_val,
                     total_params, epochs, batch_size, patience,
                     lr_max, save_dir):
        if not self.enabled:
            return
        self._add(f"Training started : {datetime.now().strftime('%Y-%m-%d  %H:%M:%S')}",
                  bold=True, size=10)
        self._add("")
        self._add(f"GPU  : {gpu_name}  ({vram_gb:.1f} GB)")
        self._add(f"Train: {n_train} patients  Val: {n_val} patients")
        self._add(f"Model: MONAI ResNet UNet 3D  ({total_params:,} params)")
        self._add("")
        self._sep()
        self._add("  MONAI 3D — cosine warm-restart run", bold=True)
        self._sep()
        self._add(f"  Epochs     : {epochs}  (warmup: {WARMUP_EPOCHS})")
        self._add(f"  Batch size : {batch_size}  (BF16 autocast)")
        self._add(f"  Patch size : {PATCH_SIZE}")
        self._add(f"  Patience   : {patience}")
        self._add(f"  LR max     : {lr_max:.0e}  (cosine restarts T0={COSINE_T0})")
        self._add(f"  Loss       : Tversky(α=0.3,β=0.7) + Focal(γ=2,α=0.95)")
        self._add(f"  Sampler    : new-dataset 2× weight")
        self._add(f"  Val thresh : sweep {VAL_THRESHOLDS}")
        self._add(f"  Save dir   : {save_dir}")
        self._sep()
        self._add("")
        self.save()

    def write_epoch(self, epoch_num, total_epochs, phase, lr,
                    train_loss, train_recall, train_prec, train_acc,
                    val_loss, val_dice, val_prec, val_rec, val_acc,
                    best_thresh, save_messages, no_improve_count):
        if not self.enabled:
            return
        self._add("")
        self._add(f"Epoch {epoch_num}/{total_epochs}  [{phase}]  LR: {lr:.2e}",
                  bold=True, color=(0, 70, 127))
        self._add(f"  Train Loss      : {train_loss:.4f}")
        self._add(f"  Train Recall    : {train_recall:.4f}")
        self._add(f"  Train Precision : {train_prec:.4f}")
        self._add(f"  Train Accuracy  : {train_acc:.4f}")
        self._add(f"  Val Loss        : {val_loss:.4f}")
        self._add(f"  Val Dice        : {val_dice:.4f}  (best-threshold sweep)")
        self._add(f"  Val Precision   : {val_prec:.4f}")
        self._add(f"  Val Recall      : {val_rec:.4f}")
        self._add(f"  Val Accuracy    : {val_acc:.4f}  (foreground only)")
        self._add(f"  Best threshold  : {best_thresh:.1f}  (modal this epoch)")
        for msg in save_messages:
            bold_m = msg.startswith(">>>") or msg.startswith("***")
            color  = (0, 100, 0) if msg.startswith("***") else (0, 0, 180)
            self._add(f"  {msg}", bold=bold_m, color=color if bold_m else None)
        if no_improve_count > 0:
            self._add(f"  No improvement ({no_improve_count}/{PATIENCE})",
                      color=(160, 80, 0))
        self.save()

    def write_footer(self, best_raw, best_smooth, final_loss,
                     best_loss, optimal_thresh, epoch_save_dir):
        if not self.enabled:
            return
        self._add("")
        self._sep()
        self._add("  MONAI 3D training complete", bold=True)
        self._sep()
        self._add(f"  Best raw Dice      : {best_raw:.4f}  → unet3d_best_raw.pth",
                  bold=True)
        self._add(f"  Best smoothed Dice : {best_smooth:.4f}  → unet3d_best.pth")
        self._add(f"  Final val loss     : {final_loss:.4f}")
        self._add(f"  Best val loss      : {best_loss:.4f}")
        self._add(f"  Optimal threshold  : {optimal_thresh:.1f}")
        self._add(f"  Epoch saves        : {epoch_save_dir}/")
        self._add("")
        self._add(f"Training finished : {datetime.now().strftime('%Y-%m-%d  %H:%M:%S')}",
                  bold=True)
        self.save()
        print(f"\n  Training log saved → {self.save_path}")


# ------------------------------------------------------------------ #
#  COLLAPSE CHECK                                                      #
# ------------------------------------------------------------------ #

def check_output_distribution(model, loader, device, epoch):
    if (epoch + 1) % 5 != 0:
        return
    model.eval()
    with torch.no_grad():
        batch  = next(iter(loader))
        images = batch["image"].to(device)
        if hasattr(images, 'as_tensor'):
            images = images.as_tensor()
        with autocast("cuda", dtype=torch.bfloat16):
            logits = model(images[:1])
        if hasattr(logits, 'as_tensor'):
            logits = logits.as_tensor()
        probs  = torch.sigmoid(logits.float())
        p_min  = probs.min().item()
        p_max  = probs.max().item()
        p_mean = probs.mean().item()
        p_high = (probs > 0.5).float().mean().item()
        if p_min < 0.05:
            status = "OK — bias working correctly"
        elif p_min > 0.20:
            status = "WARNING — still in collapse"
        else:
            status = "RECOVERING"
        print(f"\n  [Collapse ep{epoch+1}]"
              f"  min={p_min:.4f}  max={p_max:.4f}"
              f"  mean={p_mean:.4f}  frac>0.5={p_high:.4f}"
              f"  → {status}")
    model.train()


# ------------------------------------------------------------------ #
#  VALIDATION                                                          #
# ------------------------------------------------------------------ #

def validate(model, val_loader, device):
    model.eval()
    val_dices       = []
    val_losses      = []
    best_thresholds = []
    total_tp = total_fp = total_fn = total_tn = 0.0

    val_bar = tqdm(
        val_loader, desc="  Validation", leave=True, ncols=115,
        bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}] {postfix}"
    )

    with torch.no_grad():
        for batch in val_bar:
            images = batch["image"].to(device)
            labels = batch["label"].to(device)
            if hasattr(images, 'as_tensor'):
                images = images.as_tensor()
            if hasattr(labels, 'as_tensor'):
                labels = labels.as_tensor()

            with autocast("cuda", dtype=torch.bfloat16):
                preds_logits = sliding_window_inference(
                    inputs=images, roi_size=PATCH_SIZE,
                    sw_batch_size=4, predictor=model,
                    overlap=SW_OVERLAP, mode=SW_MODE,
                )
            if hasattr(preds_logits, 'as_tensor'):
                preds_logits = preds_logits.as_tensor()

            val_loss = combined_loss(preds_logits.float(), labels.float())
            val_losses.append(val_loss.item())

            prob_map = torch.sigmoid(preds_logits.float())
            gt_np    = (labels > 0).squeeze().cpu().numpy().astype(bool)

            vol_dice    = 0.0
            best_thresh = VAL_THRESHOLDS[0]

            if gt_np.sum() > 0:
                for thresh in VAL_THRESHOLDS:
                    pred_t = (prob_map > thresh).squeeze().cpu().numpy().astype(bool)
                    d      = compute_dice_np(pred_t, gt_np)
                    if d > vol_dice:
                        vol_dice    = d
                        best_thresh = thresh
                val_dices.append(vol_dice)
                best_thresholds.append(best_thresh)

            pred_bin = (prob_map > best_thresh).float()
            p = pred_bin
            t = (labels > 0).float()
            tp = (p * t).sum().item()
            fp = (p * (1 - t)).sum().item()
            fn = ((1 - p) * t).sum().item()
            tn = ((1 - p) * (1 - t)).sum().item()
            total_tp += tp;  total_fp += fp
            total_fn += fn;  total_tn += tn

            cur_prec = tp / (tp + fp + 1e-6)
            cur_rec  = tp / (tp + fn + 1e-6)
            foreground = tp + fp + fn
            cur_acc  = tp / (foreground + 1e-6) if foreground > 0 else 0.0

            val_bar.set_postfix(
                Dice=f"{vol_dice:.4f}",
                Loss=f"{val_loss.item():.4f}",
                Prec=f"{cur_prec:.4f}",
                Rec=f"{cur_rec:.4f}",
                BestT=f"{best_thresh:.1f}",
            )

    mean_dice     = float(np.mean(val_dices))  if val_dices  else 0.0
    mean_val_loss = float(np.mean(val_losses)) if val_losses else 0.0
    precision     = total_tp / (total_tp + total_fp + 1e-6)
    recall        = total_tp / (total_tp + total_fn + 1e-6)
    foreground_t  = total_tp + total_fp + total_fn
    accuracy      = total_tp / (foreground_t + 1e-6) if foreground_t > 0 else 0.0
    modal_thresh  = Counter(best_thresholds).most_common(1)[0][0] \
                    if best_thresholds else VAL_THRESHOLDS[0]

    return {
        "loss"          : mean_val_loss,
        "dice"          : mean_dice,
        "precision"     : precision,
        "recall"        : recall,
        "accuracy"      : accuracy,
        "best_threshold": modal_thresh,
    }


# ------------------------------------------------------------------ #
#  TRAINING LOOP                                                       #
# ------------------------------------------------------------------ #

def ordinal(n):
    suffix = {1:"st",2:"nd",3:"rd"}.get(
        n % 10 if n % 100 not in (11,12,13) else 0, "th")
    return f"{n}{suffix}"


def main():
    assert torch.cuda.is_available(), "CUDA GPU not found."
    device = torch.device("cuda")
    torch.cuda.set_device(0)

    gpu_name = torch.cuda.get_device_name(0)
    vram_gb  = torch.cuda.get_device_properties(0).total_memory / 1e9

    print(f"Using device: cuda")
    print(f"  GPU  : {gpu_name}")
    print(f"  VRAM : {vram_gb:.1f} GB")

    torch.set_num_threads(2)
    torch.set_num_interop_threads(1)
    torch.backends.cudnn.benchmark        = True
    torch.backends.cuda.matmul.allow_tf32 = True

    # ---- Data — weighted sampler for combined dataset ----
    print()
    train_loader, val_loader, _ = build_weighted_loaders(
        preprocessed_root   = PREPROCESSED_ROOT,
        batch_size          = BATCH_SIZE,
        num_workers         = 0,
        new_dataset_weight  = NEW_DATASET_WEIGHT,
    )

    # ---- Model ----
    print()
    model = get_model(device)
    total_params = sum(p.numel() for p in model.parameters())

    # ---- Optimiser — AdamW ----
    # BF16 autocast does NOT need GradScaler (sufficient dynamic range)
    optimizer = optim.AdamW(model.parameters(), lr=LR_START, weight_decay=1e-5)

    # ---- LR schedule: linear warmup → cosine warm restarts ----
    #
    # Phase 1: linear warmup from LR_START to LR_MAX over WARMUP_EPOCHS.
    # Phase 2: CosineAnnealingWarmRestarts — restarts every COSINE_T0
    #          epochs. Each restart resets LR to LR_MAX and decays again.
    #          This lets the model escape the plateau seen from epoch 65+
    #          in the previous single-cosine run.
    #
    # Implementation: LambdaLR for warmup, then hand off to CAWR.
    # We drive both via a single custom lambda that wraps both phases.

    def lr_lambda(epoch):
        """Returns multiplier on top of LR_START."""
        if epoch < WARMUP_EPOCHS:
            # Linear warmup: LR_START → LR_MAX
            return (epoch + 1) / WARMUP_EPOCHS * (LR_MAX / LR_START)
        else:
            # Cosine warm restart — compute position within current cycle
            ep_after_warmup = epoch - WARMUP_EPOCHS
            cycle_ep = ep_after_warmup % COSINE_T0        # position in cycle
            cos_val  = 0.5 * (1 + np.cos(np.pi * cycle_ep / COSINE_T0))
            # LR swings from LR_MAX down to LR_MAX/100 then restarts
            lr_min   = LR_MAX / 100
            lr_val   = lr_min + (LR_MAX - lr_min) * cos_val
            return lr_val / LR_START

    scheduler = optim.lr_scheduler.LambdaLR(optimizer, lr_lambda=lr_lambda)

    # ---- Training logger ----
    log_path = os.path.join(SAVE_DIR, "training_log_monai3d.docx")
    logger   = TrainingLogger(log_path)
    logger.write_header(
        gpu_name=gpu_name, vram_gb=vram_gb,
        n_train=len(train_loader.dataset),
        n_val=len(val_loader.dataset),
        total_params=total_params,
        epochs=EPOCHS, batch_size=BATCH_SIZE,
        patience=PATIENCE, lr_max=LR_MAX, save_dir=SAVE_DIR,
    )

    # ---- Training state ----
    best_raw_dice      = 0.0
    best_smoothed_dice = 0.0
    counter            = 0

    history = {
        "train_loss"    : [], "train_recall"   : [],
        "val_loss"      : [], "val_dice"       : [],
        "val_precision" : [], "val_recall"     : [],
        "val_accuracy"  : [], "val_best_thresh": [],
        "lr"            : [],
    }

    print()
    print("=" * 65)
    print("  MONAI 3D training — cosine warm-restart + weighted sampler")
    print(f"  Epochs       : {EPOCHS}  (warmup: {WARMUP_EPOCHS})")
    print(f"  Batch size   : {BATCH_SIZE}  (BF16 autocast)")
    print(f"  Patch size   : {PATCH_SIZE}")
    print(f"  Patience     : {PATIENCE}")
    print(f"  Loss         : Tversky(α=0.3,β=0.7) + Focal(γ=2,α=0.95)")
    print(f"  LR schedule  : warmup {WARMUP_EPOCHS} ep → cosine restarts T0={COSINE_T0}")
    print(f"  LR max/min   : {LR_MAX:.0e} / {LR_MAX/100:.0e}")
    print(f"  New-ds weight: {NEW_DATASET_WEIGHT}× (patients ID>{100})")
    print(f"  Val thresh   : sweep {VAL_THRESHOLDS}")
    print(f"  Save dir     : {SAVE_DIR}")
    print("=" * 65)
    print("  NOTE: Delete old .pth files before restarting:")
    print("    del models\\segmentation_3d\\*.pth")
    print("    del models\\segmentation_3d\\epochs\\*.pth")
    print("=" * 65)

    for epoch in range(EPOCHS):
        current_lr = optimizer.param_groups[0]["lr"]
        phase = "WARMUP" if epoch < WARMUP_EPOCHS else "TRAIN"

        # Annotate restart epochs
        if epoch >= WARMUP_EPOCHS:
            ep_after_warmup = epoch - WARMUP_EPOCHS
            if ep_after_warmup % COSINE_T0 == 0 and ep_after_warmup > 0:
                print(f"\n  [LR RESTART at epoch {epoch+1}] → LR reset to {LR_MAX:.0e}")

        print(f"\nEpoch {epoch+1}/{EPOCHS}  [{phase}]  LR: {current_lr:.2e}")

        # ============================================================
        # TRAIN
        # ============================================================
        model.train()
        train_loss = 0.0
        tp = fp = fn = tn = 0.0

        train_bar = tqdm(
            train_loader, desc="  Training ",
            leave=True, ncols=115,
            bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}] {postfix}"
        )

        for batch in train_bar:
            images = batch["image"].to(device)
            masks  = batch["label"].to(device)

            # Strip MetaTensor wrappers
            if hasattr(images, 'as_tensor'):
                images = images.as_tensor()
            if hasattr(masks, 'as_tensor'):
                masks = masks.as_tensor()

            masks = masks.float()
            if masks.ndim == 4:
                masks = masks.unsqueeze(1)

            optimizer.zero_grad()

            with autocast("cuda", dtype=torch.bfloat16):
                preds = model(images)
                if hasattr(preds, 'as_tensor'):
                    preds = preds.as_tensor()
                loss  = combined_loss(preds.float(), masks)

            # BF16: backward directly, no GradScaler needed
            loss.backward()
            torch.nn.utils.clip_grad_norm_(model.parameters(), max_norm=1.0)
            optimizer.step()

            train_loss += loss.item()

            with torch.no_grad():
                grad_norm = sum(
                    p.grad.norm().item() ** 2
                    for p in model.parameters() if p.grad is not None
                ) ** 0.5
                p_sig = torch.sigmoid(preds.float())
                p_bin = (p_sig > 0.5).float()
                tp += (p_bin * masks).sum().item()
                fp += (p_bin * (1 - masks)).sum().item()
                fn += ((1 - p_bin) * masks).sum().item()
                tn += ((1 - p_bin) * (1 - masks)).sum().item()

            train_bar.set_postfix(
                loss=f"{loss.item():.4f}",
                gnorm=f"{grad_norm:.3f}"
            )

        avg_train_loss = train_loss / len(train_loader)
        train_recall   = tp / (tp + fn + 1e-6)
        train_prec     = tp / (tp + fp + 1e-6)
        train_acc      = (tp + tn) / (tp + tn + fp + fn + 1e-6)

        history["train_loss"].append(avg_train_loss)
        history["train_recall"].append(train_recall)
        history["lr"].append(current_lr)

        print(f"\n  Train Loss      : {avg_train_loss:.4f}")
        print(f"  Train Recall    : {train_recall:.4f}")
        print(f"  Train Precision : {train_prec:.4f}")
        print(f"  Train Accuracy  : {train_acc:.4f}")

        scheduler.step()
        torch.cuda.empty_cache()

        # ============================================================
        # COLLAPSE CHECK  (every 5 epochs)
        # ============================================================
        check_output_distribution(model, train_loader, device, epoch)

        # ============================================================
        # VALIDATION
        # ============================================================
        print()
        val_metrics = validate(model, val_loader, device)

        avg_val_loss = val_metrics["loss"]
        avg_val_dice = val_metrics["dice"]
        prec         = val_metrics["precision"]
        rec          = val_metrics["recall"]
        acc          = val_metrics["accuracy"]
        best_thresh  = val_metrics["best_threshold"]

        history["val_loss"].append(avg_val_loss)
        history["val_dice"].append(avg_val_dice)
        history["val_precision"].append(prec)
        history["val_recall"].append(rec)
        history["val_accuracy"].append(acc)
        history["val_best_thresh"].append(best_thresh)

        print(f"\n  Val Loss        : {avg_val_loss:.4f}")
        print(f"  Val Dice        : {avg_val_dice:.4f}"
              f"  (best-threshold sweep, tumour volumes only)")
        print(f"  Val Precision   : {prec:.4f}")
        print(f"  Val Recall      : {rec:.4f}")
        print(f"  Val Accuracy    : {acc:.4f}  (foreground voxels only)")
        print(f"  Best threshold  : {best_thresh:.1f}"
              f"  (modal across val patients this epoch)")

        # ---- Smoothed Dice ----
        smoothed = float(np.mean(history["val_dice"][-3:])) \
                   if len(history["val_dice"]) >= 3 else avg_val_dice

        epoch_label      = ordinal(epoch + 1)
        _epoch_save_msgs = []

        # ---- Save every epoch ----
        torch.save(
            model.state_dict(),
            os.path.join(EPOCH_SAVE_DIR,
                         f"{epoch_label}_epoch_dice{avg_val_dice:.4f}.pth"))

        # ---- Save raw best ----
        if avg_val_dice > best_raw_dice:
            best_raw_dice = avg_val_dice
            torch.save(model.state_dict(),
                       os.path.join(SAVE_DIR, "unet3d_best_raw.pth"))
            torch.save(model.state_dict(),
                       os.path.join(EPOCH_SAVE_DIR,
                                    f"BEST_RAW_{epoch_label}_epoch"
                                    f"_dice{best_raw_dice:.4f}.pth"))
            msg = (f">>> Raw best saved    "
                   f"(Dice={best_raw_dice:.4f}, {epoch_label} epoch)"
                   f"  → unet3d_best_raw.pth")
            print(f"\n  {msg}")
            _epoch_save_msgs.append(msg)

        # ---- Save smoothed best ----
        if smoothed > best_smoothed_dice:
            best_smoothed_dice = smoothed
            counter = 0
            torch.save(model.state_dict(),
                       os.path.join(SAVE_DIR, "unet3d_best.pth"))
            torch.save(model.state_dict(),
                       os.path.join(EPOCH_SAVE_DIR,
                                    f"BEST_SMOOTH_{epoch_label}_epoch"
                                    f"_dice{avg_val_dice:.4f}.pth"))
            msg = (f"*** Smoothed best saved "
                   f"(smoothed Dice={best_smoothed_dice:.4f}, {epoch_label} epoch)"
                   f"  → unet3d_best.pth ***")
            print(f"  {msg}")
            _epoch_save_msgs.append(msg)
        else:
            counter += 1
            print(f"  No improvement ({counter}/{PATIENCE})")

        # ---- Word log ----
        logger.write_epoch(
            epoch_num=epoch+1, total_epochs=EPOCHS, phase=phase, lr=current_lr,
            train_loss=avg_train_loss, train_recall=train_recall,
            train_prec=train_prec, train_acc=train_acc,
            val_loss=avg_val_loss, val_dice=avg_val_dice,
            val_prec=prec, val_rec=rec, val_acc=acc,
            best_thresh=best_thresh,
            save_messages=_epoch_save_msgs,
            no_improve_count=counter,
        )

        # ---- Early stopping ----
        if counter >= PATIENCE:
            print(f"\nEarly stopping triggered at epoch {epoch+1}.")
            break

        np.save("training_history_3d.npy", history)

    print()
    print("=" * 65)
    print("  Training complete")
    print("=" * 65)
    print(f"  Best raw Dice      : {best_raw_dice:.4f}  → unet3d_best_raw.pth")
    print(f"  Best smoothed Dice : {best_smoothed_dice:.4f}  → unet3d_best.pth")
    if history["val_loss"]:
        print(f"  Final val loss     : {history['val_loss'][-1]:.4f}")
        print(f"  Best val loss      : {min(history['val_loss']):.4f}")
    if history["val_best_thresh"]:
        modal = Counter(history["val_best_thresh"]).most_common(1)[0][0]
        print(f"  Optimal threshold  : {modal:.1f}"
              f"  (most common across all epochs)")
    print(f"  Per-epoch saves    : {EPOCH_SAVE_DIR}/")
    print()
    print("  Next steps:")
    print("    python src/segmentation_3d/predict_3d.py")
    print("    python src/segmentation_3d/evaluate_3d.py")

    modal_thresh = Counter(history["val_best_thresh"]).most_common(1)[0][0] \
                   if history["val_best_thresh"] else VAL_THRESHOLDS[0]
    logger.write_footer(
        best_raw=best_raw_dice, best_smooth=best_smoothed_dice,
        final_loss=history["val_loss"][-1] if history["val_loss"] else 0.0,
        best_loss=min(history["val_loss"]) if history["val_loss"] else 0.0,
        optimal_thresh=modal_thresh,
        epoch_save_dir=EPOCH_SAVE_DIR,
    )


if __name__ == "__main__":
    main()
