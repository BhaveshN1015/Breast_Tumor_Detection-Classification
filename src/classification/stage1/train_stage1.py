"""
train_stage1.py
================
Stage 1 — Tumor Detection Classifier (binary: negative=0 / positive=1)

CHANGES vs previous version:
  ✓ Word .docx log updated after EVERY epoch (not just at the end)
      — crash-safe: log is complete up to last finished epoch
      — one clean summary row per epoch (no progress bar noise)
  ✓ Per-epoch model saved with ordinal name + AUC score
      (same style as train_3d.py / train_dynunet.py)
      e.g.  epochs/1st_epoch_auc0.8451.pth
  ✓ Raw best  → best_model_raw.pth   (+ labelled copy in epochs/)
  ✓ Smooth best (3-epoch rolling avg AUC) → best_model.pth
      (+ labelled copy in epochs/)
  ✓ TeeLogger REMOVED — txt log was full of progress-bar noise;
      replaced by clean Word log only
  ✓ training_log.json still saved at the end

GPU PROFILES (change ACTIVE_GPU below):
  "RTX3050"    → batch=2, fixed_size=(128,192,192), ~1.6 GB VRAM
  "RTX2000ADA" → batch=4, fixed_size=(128,192,192), ~3.0 GB VRAM

Run:
    python src/classification/stage1/train_stage1.py

Saves to:
    models/classification_stage1/
        best_model_raw.pth          ← highest single-epoch val AUC
        best_model.pth              ← highest 3-epoch smoothed val AUC
        last_model.pth              ← checkpoint after every epoch
        training_log.json           ← full history (written at end)
        training_log_stage1.docx    ← Word log updated after EVERY epoch
        epochs/
            1st_epoch_auc0.8451.pth
            2nd_epoch_auc0.8589.pth
            ...
            BEST_RAW_2nd_epoch_auc0.8589.pth
            BEST_SMOOTH_3rd_epoch_auc0.8530.pth
"""

import os
import sys
import json
import time
import datetime
import numpy as np
import torch
import torch.nn as nn
from torch.amp import GradScaler, autocast
from torch.optim import AdamW
from torch.optim.lr_scheduler import CosineAnnealingWarmRestarts
from tqdm import tqdm

PROJECT_ROOT = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "..")
)
sys.path.insert(0, PROJECT_ROOT)

from src.classification.stage1.dataset_stage1 import get_dataloaders, FIXED_SIZE
from src.classification.stage1.model_stage1   import build_model

try:
    from sklearn.metrics import roc_auc_score, f1_score, confusion_matrix
    SKLEARN_OK = True
except ImportError:
    SKLEARN_OK = False
    print("[WARN] scikit-learn not found. Run: pip install scikit-learn")

# ── python-docx: auto-install if missing ────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    try:
        import subprocess
        subprocess.run(
            [sys.executable, "-m", "pip", "install", "python-docx",
             "--quiet", "--break-system-packages"],
            check=True
        )
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        DOCX_OK = True
    except Exception:
        DOCX_OK = False
        print("[WARN] python-docx unavailable — Word log will be skipped")


# ================================================================== #
#  SELECT GPU PROFILE                                                  #
# ================================================================== #

ACTIVE_GPU = "RTX3050"   # ← change to "RTX2000ADA" when using Ada

GPU_PROFILES = {
    "RTX3050": {
        "batch_size" : 4,
        "fixed_size" : (128, 192, 192),
        "num_workers": 0,
        "use_amp"    : True,
        "note"       : "RTX 3050 6GB — ~1.6 GB VRAM at batch=2",
    },
    "RTX2000ADA": {
        "batch_size" : 4,
        "fixed_size" : (128, 192, 192),
        "num_workers": 4,
        "use_amp"    : True,
        "note"       : "RTX 2000 Ada 16GB — ~3.0 GB VRAM at batch=4",
    },
}
GPU_CFG = GPU_PROFILES[ACTIVE_GPU]


# ================================================================== #
#  CONFIG                                                              #
# ================================================================== #

CFG = {
    "data_root"  : r"D:\Breast_Tumor_AI_Project\data\classification_stage1",
    "save_dir"   : r"D:\Breast_Tumor_AI_Project\models\classification_stage1",

    "pos_weight" : 2.0,
    "dropout"    : 0.5,

    "epochs"     : 80,
    "lr"         : 3e-4,
    "weight_decay": 2e-4,
    "T0"         :15,
    "T_mult"     : 1,

    "monitor"    : "val_auc",
    "patience"   : 20,
    "label_smooth_eps": 0.1,
    **GPU_CFG,
}


# ================================================================== #
#  ORDINAL HELPER  (1st, 2nd, 3rd … identical to train_3d.py)        #
# ================================================================== #

def ordinal(n: int) -> str:
    suffix = ["th", "st", "nd", "rd"] + ["th"] * 16
    return f"{n}{suffix[n % 20] if n % 20 < len(suffix) else 'th'}"


# ================================================================== #
#  METRICS                                                             #
# ================================================================== #

def compute_metrics(labels, probs, preds, prefix):
    m = {}
    acc = sum(p == l for p, l in zip(preds, labels)) / max(len(labels), 1)
    m[f"{prefix}_acc"] = round(acc, 4)

    if SKLEARN_OK and len(set(labels)) > 1:
        m[f"{prefix}_auc"]  = round(float(roc_auc_score(labels, probs)), 4)
        m[f"{prefix}_f1"]   = round(float(
            f1_score(labels, preds, zero_division=0)), 4)
        cm = confusion_matrix(labels, preds, labels=[0, 1])
        if cm.shape == (2, 2):
            tn, fp, fn, tp = cm.ravel()
            m[f"{prefix}_sens"] = round(tp/(tp+fn) if (tp+fn) > 0 else 0, 4)
            m[f"{prefix}_spec"] = round(tn/(tn+fp) if (tn+fp) > 0 else 0, 4)
            m[f"{prefix}_ppv"]  = round(tp/(tp+fp) if (tp+fp) > 0 else 0, 4)
    else:
        m[f"{prefix}_auc"] = 0.0
        m[f"{prefix}_f1"]  = 0.0
    return m


# ================================================================== #
#  WORD LOG — updated after every epoch (crash-safe)                  #
# ================================================================== #

class WordLogger:
    """
    Maintains a .docx that is updated and saved to disk after every
    epoch. If training crashes at epoch 40, you have epochs 1-39 in
    the file. Format is identical to train_3d.py / train_dynunet.py.
    """

    def __init__(self, path: str):
        self.path = path
        self.doc  = None
        if not DOCX_OK:
            return
        self.doc = Document()
        # Consolas 9pt — same as segmentation logs
        self.doc.styles["Normal"].font.name = "Consolas"
        self.doc.styles["Normal"].font.size = Pt(9)

    # ── Internal helpers ──────────────────────────────────────────

    def _p(self, text: str, bold: bool = False,
           color: tuple = None) -> None:
        """Add one paragraph in Consolas 9pt."""
        if self.doc is None:
            return
        para = self.doc.add_paragraph()
        run  = para.add_run(text)
        run.bold           = bold
        run.font.name      = "Consolas"
        run.font.size      = Pt(9)
        if color:
            run.font.color.rgb = RGBColor(*color)

    def _save(self) -> None:
        if self.doc is None:
            return
        try:
            self.doc.save(self.path)
        except Exception as exc:
            # Never crash training because of log saving
            print(f"  [WARN] Could not save Word log: {exc}")

    # ── Public API ────────────────────────────────────────────────

    def write_header(self, device, gpu_name: str, vram_gb: float,
                     train_n: int, val_n: int, test_n: int) -> None:
        """Called once before epoch 1."""
        if self.doc is None:
            return
        self._p("=" * 70, bold=True)
        self._p("  STAGE 1 TRAINING LOG — Tumor Detection Classifier",
                bold=True)
        self._p("  Model: 3D ResNet18 + CBAM Attention", bold=True)
        self._p("=" * 70, bold=True)
        self._p(f"  Started   : {datetime.datetime.now():%Y-%m-%d %H:%M:%S}")
        self._p(f"  Device    : {device}")
        if gpu_name:
            self._p(f"  GPU       : {gpu_name}  ({vram_gb:.1f} GB VRAM)")
        self._p(f"  Profile   : {ACTIVE_GPU}  — {GPU_CFG['note']}")
        self._p(f"  Dataset   : train={train_n}  val={val_n}  test={test_n}")
        self._p(f"  Batch     : {CFG['batch_size']}   "
                f"Fixed size: {CFG['fixed_size']}   "
                f"Dropout: {CFG['dropout']}")
        self._p(f"  LR        : {CFG['lr']:.0e}   "
                f"T0: {CFG['T0']}   Patience: {CFG['patience']}   "
                f"Epochs: {CFG['epochs']}")
        self._p(f"  pos_weight: {CFG['pos_weight']}   AMP: {CFG['use_amp']}")
        self._p(f"  Save dir  : {CFG['save_dir']}")
        self._p("=" * 70, bold=True)
        self._save()

    def write_epoch(self, epoch: int, elapsed: float, lr: float,
                    train_m: dict, val_m: dict,
                    save_msgs: list) -> None:
        """
        Called after every epoch. Writes a clean summary block —
        no progress-bar noise, just the final metrics.
        """
        if self.doc is None:
            return

        # Determine LR-phase tag
        t0_size = CFG["T0"]
        if epoch <= 5:
            tag = "[WARMUP]"
        else:
            cycle_ep = (epoch - 1) % t0_size
            tag = f"[Cycle ep {cycle_ep + 1}/{t0_size}]"

        ep_ord = ordinal(epoch)
        self._p(f"\nEpoch {ep_ord}/{CFG['epochs']}  "
                f"{tag}  LR: {lr:.2e}  time={elapsed:.0f}s",
                bold=True)

        # Train metrics
        self._p(
            f"  Train  loss={train_m.get('train_loss', 0):.4f}  "
            f"acc={train_m.get('train_acc', 0):.4f}  "
            f"auc={train_m.get('train_auc', 0):.4f}  "
            f"f1={train_m.get('train_f1', 0):.4f}"
        )
        # Val metrics (bold + blue — new best will stand out)
        val_line = (
            f"  Val    loss={val_m.get('val_loss', 0):.4f}  "
            f"acc={val_m.get('val_acc', 0):.4f}  "
            f"auc={val_m.get('val_auc', 0):.4f}  "
            f"f1={val_m.get('val_f1', 0):.4f}"
        )
        self._p(val_line, bold=True, color=(0, 100, 200))

        self._p(
            f"         sens={val_m.get('val_sens', 0):.4f}  "
            f"spec={val_m.get('val_spec', 0):.4f}  "
            f"ppv={val_m.get('val_ppv', 0):.4f}"
        )

        # Save messages (green for smooth best, amber for raw best)
        for msg in save_msgs:
            color = (0, 140, 0) if "SMOOTH" in msg else (180, 100, 0)
            self._p(f"  {msg}", color=color)

        self._save()   # flush to disk immediately — crash-safe

    def write_footer(self, best_raw_epoch: int, best_raw_auc: float,
                     best_smooth_auc: float, test_m: dict) -> None:
        """Called once after training completes."""
        if self.doc is None:
            return
        self._p("\n" + "=" * 70, bold=True)
        self._p("  TRAINING COMPLETE", bold=True)
        self._p(f"  Finished  : {datetime.datetime.now():%Y-%m-%d %H:%M:%S}")
        self._p(f"  Best raw epoch  : {ordinal(best_raw_epoch)}  "
                f"AUC={best_raw_auc:.4f}")
        self._p(f"  Best smooth AUC : {best_smooth_auc:.4f}")
        if test_m:
            self._p("  ── Test set results ──")
            for k, v in test_m.items():
                label = k.replace("test_", "").capitalize()
                self._p(f"    {label:20s}: "
                        f"{v:.4f}" if isinstance(v, float) else f"    {label}: {v}")
        self._p(f"  Epoch file format : {ordinal(1)}_epoch_auc0.XXXX.pth")
        self._p(f"  best_model_raw.pth  — highest single-epoch AUC")
        self._p(f"  best_model.pth      — highest 3-epoch smooth AUC")
        self._p("=" * 70, bold=True)
        self._save()
        print(f"\n  Word log → {self.path}")


# ================================================================== #
#  TRAIN ONE EPOCH                                                     #
# ================================================================== #

def train_one_epoch(model, loader, criterion, optimizer, scaler,
                    device, use_amp, epoch, total_epochs):
    model.train()
    total_loss = 0.0
    all_labels, all_probs, all_preds = [], [], []

    pbar = tqdm(
        loader,
        desc=f"  Epoch {epoch:3d}/{total_epochs} [TRAIN]",
        total=len(loader),
        ncols=80,
        leave=False,
        dynamic_ncols=False,
        file=sys.stdout,
    )

    for batch in pbar:
        images = batch["image"].to(device, non_blocking=True)
        labels = batch["label"].to(device, non_blocking=True).unsqueeze(1)

        optimizer.zero_grad(set_to_none=True)

        with autocast(device_type=device.type, enabled=use_amp):
            logits = model(images)
            loss = criterion(logits, smooth_labels(labels))

        if use_amp and device.type == "cuda":
            scaler.scale(loss).backward()
            scaler.unscale_(optimizer)
            nn.utils.clip_grad_norm_(model.parameters(), 1.0)
            scaler.step(optimizer)
            scaler.update()
        else:
            loss.backward()
            nn.utils.clip_grad_norm_(model.parameters(), 1.0)
            optimizer.step()

        total_loss += loss.item()
        probs = torch.sigmoid(logits).detach().cpu().numpy().flatten()
        preds = (probs >= 0.5).astype(int)
        labs  = labels.detach().cpu().numpy().flatten().astype(int)
        all_probs.extend(probs.tolist())
        all_preds.extend(preds.tolist())
        all_labels.extend(labs.tolist())
        pbar.set_postfix({"loss": f"{loss.item():.4f}"}, refresh=True)

    pbar.close()
    m = compute_metrics(all_labels, all_probs, all_preds, "train")
    m["train_loss"] = round(total_loss / len(loader), 5)
    return m


# ================================================================== #
#  EVALUATE                                                            #
# ================================================================== #

@torch.no_grad()
def evaluate(model, loader, criterion, device, use_amp,
             prefix="val", show_bar=True, epoch=None, total_epochs=None):
    model.eval()
    total_loss = 0.0
    all_labels, all_probs, all_preds = [], [], []

    desc = (f"  Epoch {epoch:3d}/{total_epochs} [{prefix.upper():5s}]"
            if epoch else f"  [{prefix.upper()}]")

    iter_obj = tqdm(
        loader,
        desc=desc,
        ncols=80,
        leave=False,
        dynamic_ncols=False,
        file=sys.stdout,
    ) if show_bar else loader

    for batch in iter_obj:
        images = batch["image"].to(device, non_blocking=True)
        labels = batch["label"].to(device, non_blocking=True).unsqueeze(1)

        with autocast(device_type=device.type, enabled=use_amp):
            logits = model(images)
            loss   = criterion(logits, labels)

        total_loss += loss.item()
        probs = torch.sigmoid(logits).cpu().numpy().flatten()
        preds = (probs >= 0.5).astype(int)
        labs  = labels.cpu().numpy().flatten().astype(int)
        all_probs.extend(probs.tolist())
        all_preds.extend(preds.tolist())
        all_labels.extend(labs.tolist())

    if show_bar:
        iter_obj.close()

    m = compute_metrics(all_labels, all_probs, all_preds, prefix)
    m[f"{prefix}_loss"] = round(total_loss / len(loader), 5)
    return m


# ================================================================== #
#  PRINT EPOCH SUMMARY  (terminal)                                    #
# ================================================================== #

def print_epoch_summary(epoch, total_epochs, train_m, val_m,
                         elapsed, is_best, best_score, lr):
    sep = "  " + "─" * 61
    print(sep)
    print(f"  Epoch {epoch:3d}/{total_epochs}   "
          f"time={elapsed:.0f}s   lr={lr:.2e}   "
          f"{'★ NEW BEST' if is_best else ''}")
    print()
    print(f"  {'Metric':20s}  {'Train':>9s}  {'Val':>9s}")
    print(f"  {'-'*20}  {'-'*9}  {'-'*9}")
    rows = [
        ("Loss",        "train_loss",  "val_loss"),
        ("Accuracy",    "train_acc",   "val_acc"),
        ("AUC",         "train_auc",   "val_auc"),
        ("F1 Score",    "train_f1",    "val_f1"),
        ("Sensitivity", "train_sens",  "val_sens"),
        ("Specificity", "train_spec",  "val_spec"),
    ]
    for label, tk, vk in rows:
        tv = train_m.get(tk, "—")
        vv = val_m.get(vk, "—")
        tv_s = f"{tv:.4f}" if isinstance(tv, float) else str(tv)
        vv_s = f"{vv:.4f}" if isinstance(vv, float) else str(vv)
        print(f"  {label:20s}  {tv_s:>9s}  {vv_s:>9s}")
    if is_best:
        print(f"\n  ✓ Best model saved  (val_auc={best_score:.4f})")
    print(sep)
    print()

def smooth_labels(labels, eps=0.1):
    """Smooth binary labels: 0 → 0.1, 1 → 0.9"""
    return labels * (1 - eps) + eps * 0.5

# ================================================================== #
#  MAIN                                                                #
# ================================================================== #

def main():
    os.makedirs(CFG["save_dir"], exist_ok=True)
    epochs_dir = os.path.join(CFG["save_dir"], "epochs")
    os.makedirs(epochs_dir, exist_ok=True)

    start_dt = datetime.datetime.now()

    print()
    print("=" * 65)
    print("  Stage 1 — Tumor Detection Classifier")
    print("  Model : 3D ResNet18 + CBAM Attention")
    print(f"  GPU   : {ACTIVE_GPU}  — {GPU_CFG['note']}")
    print("=" * 65)

    # ── Device ────────────────────────────────────────────────────
    device   = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    gpu_name = ""
    vram_gb  = 0.0
    print(f"\n  Device    : {device}")
    if device.type == "cuda":
        gpu_name = torch.cuda.get_device_name(0)
        vram_gb  = torch.cuda.get_device_properties(0).total_memory / 1e9
        print(f"  GPU       : {gpu_name}")
        print(f"  VRAM      : {vram_gb:.1f} GB")

    # ── Data ──────────────────────────────────────────────────────
    print(f"\n  Data root : {CFG['data_root']}")
    print(f"  Fixed size: {CFG['fixed_size']}  (D×H×W — padded+cropped)")
    print(f"  Batch size: {CFG['batch_size']}")
    print(f"  Workers   : {CFG['num_workers']}")
    print()

    train_loader, val_loader, test_loader = get_dataloaders(
        root_dir    = CFG["data_root"],
        batch_size  = CFG["batch_size"],
        num_workers = CFG["num_workers"],
        fixed_size  = CFG["fixed_size"],
    )
    train_n = len(train_loader.dataset)
    val_n   = len(val_loader.dataset)
    test_n  = len(test_loader.dataset)

    # ── Model ─────────────────────────────────────────────────────
    print()
    model = build_model(in_channels=3, dropout=CFG["dropout"])
    model = model.to(device)

    pos_weight = torch.tensor([CFG["pos_weight"]], device=device)
    criterion = nn.BCEWithLogitsLoss(pos_weight=pos_weight,reduction="mean")

    print(f"  Loss      : BCEWithLogitsLoss  pos_weight={CFG['pos_weight']}")

    optimizer = AdamW(model.parameters(),
                      lr=CFG["lr"], weight_decay=CFG["weight_decay"])
    scheduler = CosineAnnealingWarmRestarts(
        optimizer, T_0=CFG["T0"], T_mult=CFG["T_mult"])
    scaler    = GradScaler("cuda",
                           enabled=(CFG["use_amp"] and device.type == "cuda"))

    print(f"  Epochs    : {CFG['epochs']}")
    print(f"  LR        : {CFG['lr']}")
    print(f"  AMP       : {CFG['use_amp']}")
    print(f"  Patience  : {CFG['patience']}")
    print()

    # ── Word logger ───────────────────────────────────────────────
    log_docx_path = os.path.join(CFG["save_dir"], "training_log_stage1.docx")
    wlog = WordLogger(log_docx_path)
    wlog.write_header(device, gpu_name, vram_gb, train_n, val_n, test_n)

    # ── Training state ────────────────────────────────────────────
    best_raw_auc    = -1.0
    best_smooth_auc = -1.0
    smooth_window   = []          # last 3 val AUC values
    best_raw_epoch  = 1
    patience_cnt    = 0
    history         = []

    for epoch in range(1, CFG["epochs"] + 1):
        t0 = time.time()

        # ── Train ──────────────────────────────────────────────────
        train_m = train_one_epoch(
            model, train_loader, criterion, optimizer, scaler,
            device, CFG["use_amp"], epoch, CFG["epochs"])

        # ── Validate ───────────────────────────────────────────────
        val_m = evaluate(
            model, val_loader, criterion, device, CFG["use_amp"],
            prefix="val", show_bar=True,
            epoch=epoch, total_epochs=CFG["epochs"])

        scheduler.step()
        elapsed = time.time() - t0
        lr_now  = optimizer.param_groups[0]["lr"]

        row = {"epoch": epoch, **train_m, **val_m, "lr": round(lr_now, 7)}
        history.append(row)

        # ── Every-epoch ordinal checkpoint ─────────────────────────
        # e.g.  epochs/1st_epoch_auc0.8451.pth
        cur_auc   = val_m.get("val_auc", 0.0)
        ep_name   = f"{ordinal(epoch)}_epoch_auc{cur_auc:.4f}.pth"
        ep_path   = os.path.join(epochs_dir, ep_name)
        torch.save({
            "epoch"      : epoch,
            "model_state": model.state_dict(),
            "val_auc"    : cur_auc,
            "cfg"        : CFG,
        }, ep_path)

        save_msgs = []

        # ── Raw best ───────────────────────────────────────────────
        is_best = cur_auc > best_raw_auc
        if is_best:
            best_raw_auc   = cur_auc
            best_raw_epoch = epoch
            patience_cnt   = 0

            # Main save (used by evaluate_stage1.py)
            torch.save({
                "epoch"      : epoch,
                "model_state": model.state_dict(),
                "optimizer"  : optimizer.state_dict(),
                "scheduler"  : scheduler.state_dict(),
                "best_score" : best_raw_auc,
                "cfg"        : CFG,
            }, os.path.join(CFG["save_dir"], "best_model_raw.pth"))

            # Labelled copy in epochs/
            raw_lbl = os.path.join(
                epochs_dir,
                f"BEST_RAW_{ordinal(epoch)}_epoch_auc{cur_auc:.4f}.pth")
            torch.save(model.state_dict(), raw_lbl)

            msg = (f">>> Raw best saved  — AUC={cur_auc:.4f}  "
                   f"Acc={val_m.get('val_acc', 0):.4f}  "
                   f"(epoch {ordinal(epoch)})")
            print(f"  {msg}")
            save_msgs.append(msg)
        else:
            patience_cnt += 1

        # ── Smooth best  (3-epoch rolling mean AUC) ────────────────
        smooth_window.append(cur_auc)
        if len(smooth_window) > 3:
            smooth_window.pop(0)
        smooth_auc = float(np.mean(smooth_window))

        if smooth_auc > best_smooth_auc:
            best_smooth_auc = smooth_auc

            # Main save (default checkpoint used by most downstream code)
            torch.save({
                "epoch"      : epoch,
                "model_state": model.state_dict(),
                "optimizer"  : optimizer.state_dict(),
                "scheduler"  : scheduler.state_dict(),
                "best_score" : best_smooth_auc,
                "cfg"        : CFG,
            }, os.path.join(CFG["save_dir"], "best_model.pth"))

            # Labelled copy in epochs/
            smooth_lbl = os.path.join(
                epochs_dir,
                f"BEST_SMOOTH_{ordinal(epoch)}_epoch_auc{smooth_auc:.4f}.pth")
            torch.save(model.state_dict(), smooth_lbl)

            msg = (f"*** Smooth best saved — "
                   f"smooth_AUC={smooth_auc:.4f}  "
                   f"(epoch {ordinal(epoch)})")
            print(f"  {msg}")
            save_msgs.append(msg)

        # ── last_model (always overwritten) ────────────────────────
        torch.save({
            "epoch"      : epoch,
            "model_state": model.state_dict(),
            "optimizer"  : optimizer.state_dict(),
            "scheduler"  : scheduler.state_dict(),
        }, os.path.join(CFG["save_dir"], "last_model.pth"))

        # ── Terminal summary ────────────────────────────────────────
        print_epoch_summary(
            epoch, CFG["epochs"], train_m, val_m,
            elapsed, is_best, best_raw_auc, lr_now)

        # ── Word log update — written to disk immediately ───────────
        wlog.write_epoch(epoch, elapsed, lr_now,
                         train_m, val_m, save_msgs)

        # ── Early stopping ─────────────────────────────────────────
        if patience_cnt >= CFG["patience"]:
            print(f"  Early stopping triggered at epoch {epoch}")
            print(f"  (no improvement for {CFG['patience']} epochs)\n")
            break

        # ── Periodic VRAM cache clear ──────────────────────────────
        if (epoch % 10) == 0:
            torch.cuda.empty_cache()

    # ================================================================ #
    #  FINAL TEST EVALUATION                                           #
    # ================================================================ #

    print("=" * 65)
    print("  Final TEST evaluation (best_model_raw.pth)")
    print("=" * 65)

    raw_path = os.path.join(CFG["save_dir"], "best_model_raw.pth")
    ckpt = torch.load(raw_path, map_location=device, weights_only=False)
    model.load_state_dict(ckpt["model_state"])

    test_m = evaluate(model, test_loader, criterion, device,
                      CFG["use_amp"], prefix="test", show_bar=True)

    print()
    print(f"  {'Metric':20s}  {'Value':>9s}")
    print(f"  {'-'*20}  {'-'*9}")
    for k, v in test_m.items():
        label = k.replace("test_", "").capitalize()
        print(f"  {label:20s}  "
              f"{f'{v:.4f}' if isinstance(v, float) else str(v):>9s}")

    # ── Save JSON log ─────────────────────────────────────────────
    log_json = os.path.join(CFG["save_dir"], "training_log.json")
    with open(log_json, "w") as f:
        json.dump({"cfg": CFG, "history": history,
                   "test_results": test_m},
                  f, indent=2, default=str)

    # ── Finish Word log ───────────────────────────────────────────
    wlog.write_footer(best_raw_epoch, best_raw_auc,
                      best_smooth_auc, test_m)

    # ── Completion summary ────────────────────────────────────────
    print(f"\n{'='*65}")
    print(f"  TRAINING COMPLETE")
    print(f"  Best raw   AUC : {best_raw_auc:.4f}  "
          f"(epoch {ordinal(best_raw_epoch)})")
    print(f"  Best smooth AUC: {best_smooth_auc:.4f}")
    print(f"  Epoch files    : {ordinal(1)}_epoch_auc0.XXXX.pth")
    print(f"  best_model_raw.pth  — single best epoch")
    print(f"  best_model.pth      — smoothed best")
    print(f"  Word log       : training_log_stage1.docx  (complete)")
    print(f"  JSON log       : {log_json}")
    print(f"{'='*65}\n")


if __name__ == "__main__":
    main()